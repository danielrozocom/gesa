import os
import re
import copy
import glob
import json
import time
import zipfile
import tempfile
import shutil
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

CM_TO_PT = 28.3465


def _open_doc(path):
    """Abre un .docx con python-docx."""
    return Document(path)


# ─── CONFIGURACIÓN DE GRADOS Y NIVELES ─────────────────────────

GRADES_INFO = {
    "Prejardín": {"level": "Preescolar", "multicourse": False},
    "Jardín": {"level": "Preescolar", "multicourse": False},
    "Transición": {"level": "Preescolar", "multicourse": False},
    "1°": {"level": "Básica Primaria", "multicourse": False},
    "2°": {"level": "Básica Primaria", "multicourse": False},
    "3°": {"level": "Básica Primaria", "multicourse": False},
    "4°": {"level": "Básica Primaria", "multicourse": False},
    "5°": {"level": "Básica Primaria", "multicourse": False},
    "6°": {"level": "Básica Secundaria", "multicourse": True},
    "7°": {"level": "Básica Secundaria", "multicourse": True},
    "8°": {"level": "Básica Secundaria", "multicourse": True},
    "9°": {"level": "Básica Secundaria", "multicourse": True},
    "10°": {"level": "Media Vocacional", "multicourse": True},
    "11°": {"level": "Media Vocacional", "multicourse": True}
}

MONTHS_LIST = ['ENE', 'FEB', 'MAR', 'ABR', 'MAY', 'JUN', 'JUL', 'AGO', 'SEP', 'OCT', 'NOV', 'DIC']


SHORTCODES = {
    "grade": "Grado (ej. 2°)",
    "period": "Periodo (ej. P3)",
    "session": "Código de subsesión (ej. S1.1)",
    "year": "Año (ej. 2026)",
    "level": "Nivel educativo (ej. Media Vocacional)",
}


def expand_template(template, context):
    return re.sub(r'\{(\w+)\}', lambda m: str(context.get(m.group(1), m.group(0))), template)


# ─── helpers ───────────────────────────────────────────────────

CM_TO_PT = 72.0 / 2.54  # 1 cm = 28.3465 pt

# Patrón de caracteres inválidos en XML 1.0 (se eliminan del texto antes de insertar en XML)
INVALID_XML_CHARS = re.compile(
    '[\x00-\x08\x0B\x0C\x0E-\x1F\uFFFE\uFFFF]'
)

def sanitize_xml_text(text):
    """Elimina caracteres que XML 1.0 no permite."""
    if text is None:
        return None
    return INVALID_XML_CHARS.sub('', text)

def reorder_sectPr(sectPr):
    """Reordena los hijos de w:sectPr según el esquema OpenXML (ECMA-376)."""
    if sectPr is None:
        return
    order = [
        'headerReference', 'footerReference',
        'footnotePr', 'endnotePr', 'type', 'pgSz', 'pgMar', 'paperSrc',
        'pgBorders', 'lnNumType', 'cols', 'formProt', 'vAlign', 'noEndnote',
        'titlePg', 'textDirection', 'bidi', 'rtlGutter', 'docGrid',
        'printerSettings', 'sectPrChange'
    ]
    tag_map = {tag: i for i, tag in enumerate(order)}
    children = list(sectPr)
    if not children:
        return

    def key_fn(elem):
        t = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
        return tag_map.get(t, 999)

    sorted_children = sorted(children, key=key_fn)
    if sorted_children != children:
        for elem in children:
            sectPr.remove(elem)
        for elem in sorted_children:
            sectPr.append(elem)

def sanitize_document_xml(doc):
    """Limpia todo el contenido textual del documento: elimina caracteres XML inválidos
    y asegura estructura correcta del body y sectPr según el esquema OpenXML."""
    if doc is None:
        return
    wns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    body = doc.element.body

    # Limpiar texto en todos los w:t
    for t in body.iter(f'{{{wns}}}t'):
        if t.text:
            t.text = sanitize_xml_text(t.text)
        if t.tail:
            t.tail = sanitize_xml_text(t.tail)

    # Remover w:sectPr de cualquier padre que no sea body (pPr, customXml, etc.)
    for sectPr in list(body.xpath('.//w:sectPr')):
        parent = sectPr.getparent()
        if parent is not None and parent != body:
            parent.remove(sectPr)

    # Asegurar exactamente 1 w:sectPr al final del body
    body_sectPrs = body.xpath('w:sectPr')
    if len(body_sectPrs) > 1:
        keep = body_sectPrs[-1]
        for sp in body_sectPrs[:-1]:
            body.remove(sp)

    # Reordenar hijos del body sectPr según esquema OpenXML
    body_sectPr = body.find(f'{{{wns}}}sectPr')
    if body_sectPr is not None:
        reorder_sectPr(body_sectPr)

    # Reordenar sectPr dentro de otros elementos (paragrafos con sectPr removidos arriba)
    # También cualquier sectPr en headers, footers etc.
    for sectPr in body.xpath('.//w:sectPr'):
        reorder_sectPr(sectPr)

    # Limpiar también headers, footers
    for sec in doc.sections:
        for h in [sec.header, sec.first_page_header, sec.footer]:
            if h is not None:
                for t in h._element.iter(f'{{{wns}}}t'):
                    if t.text:
                        t.text = sanitize_xml_text(t.text)
                    if t.tail:
                        t.tail = sanitize_xml_text(t.tail)


def get_document_default_font(doc):
    """Detect default font name from document styles/defaults if present."""
    if doc is None:
        return None
    try:
        if hasattr(doc, 'styles') and 'Normal' in doc.styles:
            normal = doc.styles['Normal']
            if normal.font and normal.font.name:
                return normal.font.name
        styles_part = getattr(doc.part, 'styles_part', None)
        if styles_part is not None:
            rPr = styles_part.element.xpath('.//w:docDefaults/w:rPrDefault/w:rPr/w:rFonts')
            if rPr:
                ascii_font = rPr[0].get(qn('w:ascii')) or rPr[0].get(qn('w:hAnsi'))
                if ascii_font:
                    return ascii_font
    except Exception:
        pass
    return None


def freeze_subdocument_fonts(doc):
    """Explicitly preserve/freeze subdocument run fonts and sizes so they do not inherit master template styles when merged."""
    if doc is None:
        return
    wns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    sub_font = get_document_default_font(doc)
    sub_size_val = None
    try:
        if hasattr(doc, 'styles') and 'Normal' in doc.styles:
            n_font = doc.styles['Normal'].font
            if n_font and n_font.size:
                sub_size_val = str(int(n_font.size.pt * 2))
    except Exception:
        pass

    for r_elem in doc.element.body.xpath('.//w:r'):
        rPr = r_elem.find(f'{{{wns}}}rPr')
        if rPr is None:
            rPr = parse_xml(f'<w:rPr {nsdecls("w")}/>')
            r_elem.insert(0, rPr)

        if sub_font:
            rFonts = rPr.find(f'{{{wns}}}rFonts')
            if rFonts is None:
                rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{sub_font}" w:hAnsi="{sub_font}" w:cs="{sub_font}"/>')
                rPr.append(rFonts)

        if sub_size_val:
            sz = rPr.find(f'{{{wns}}}sz')
            if sz is None:
                sz = parse_xml(f'<w:sz {nsdecls("w")} w:val="{sub_size_val}"/>')
                rPr.append(sz)
            szCs = rPr.find(f'{{{wns}}}szCs')
            if szCs is None:
                szCs = parse_xml(f'<w:szCs {nsdecls("w")} w:val="{sub_size_val}"/>')
                rPr.append(szCs)


def reorder_pPr(pPr):
    """Reorders children of a w:pPr element according to strict OpenXML schema specification,
    and removes any duplicate children of the same tag."""
    if pPr is None:
        return
    order = [
        'pStyle', 'keepNext', 'keepLines', 'pageBreakBefore', 'framePr',
        'widowControl', 'numPr', 'suppressLineNumbers', 'pBdr', 'shd',
        'tabs', 'suppressAutoHyphens', 'kinsoku', 'wordWrap', 'overflowPunct',
        'topLinePunct', 'autoSpaceDE', 'autoSpaceDN', 'bidi', 'adjustRightInd',
        'snapToGrid', 'spacing', 'ind', 'contextualSpacing', 'mirrorIndents',
        'suppressOverlap', 'jc', 'textDirection', 'textAlignment',
        'textboxTight', 'outlineLvl', 'rPr', 'sectPr', 'pPrChange'
    ]
    tag_map = {tag: i for i, tag in enumerate(order)}

    children = list(pPr)
    if not children:
        return

    singletons = {'pStyle', 'numPr', 'spacing', 'ind', 'jc', 'rPr', 'sectPr', 'outlineLvl'}
    seen = {}
    for elem in children:
        tag_local = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
        if tag_local in singletons:
            seen[tag_local] = elem

    for elem in children:
        tag_local = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
        if tag_local in singletons and seen.get(tag_local) != elem:
            pPr.remove(elem)

    remaining = list(pPr)
    def key_fn(elem):
        t = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
        return tag_map.get(t, 999)

    sorted_children = sorted(remaining, key=key_fn)
    if sorted_children != remaining:
        for elem in remaining:
            pPr.remove(elem)
        for elem in sorted_children:
            pPr.append(elem)


# Orden estricto de hijos según ECMA-376 para elementos que Word valida.
# Ordenar alfabéticamente produce esquemas inválidos (ej. <w:b/> antes de
# <w:rFonts/>) que Word rechaza con "contenido ilegible".
_SCHEMA_ORDERS = {
    'rPr': [
        'rStyle', 'rFonts', 'b', 'bCs', 'i', 'iCs', 'caps', 'smallCaps',
        'strike', 'dstrike', 'outline', 'shadow', 'emboss', 'imprint',
        'noProof', 'snapToGrid', 'vanish', 'webHidden', 'color', 'spacing',
        'w', 'kern', 'position', 'sz', 'szCs', 'highlight', 'u', 'effect',
        'bdr', 'shd', 'fitText', 'vertAlign', 'rtl', 'cs', 'em', 'lang',
        'eastAsianLayout', 'specVanish', 'oMath',
    ],
    'tblPr': [
        'tblStyle', 'tblpPr', 'tblOverlap', 'bidiVisual', 'tblStyleRowBandSize',
        'tblStyleColBandSize', 'tblW', 'jc', 'tblCellSpacing', 'tblInd',
        'tblBorders', 'shd', 'tblLayout', 'tblCellMar', 'tblLook',
        'tblCaption', 'tblDescription',
    ],
    'tcPr': [
        'cnfStyle', 'tcW', 'gridSpan', 'hMerge', 'vMerge', 'tcBorders', 'shd',
        'noWrap', 'tcMar', 'textDirection', 'tcFitText', 'vAlign', 'hideMark',
        'headers', 'cellIns', 'cellDel', 'cellMerge', 'tcPrChange',
    ],
    'trPr': [
        'cnfStyle', 'divId', 'gridBefore', 'gridAfter', 'wBefore', 'wAfter',
        'cantSplit', 'trHeight', 'tblHeader', 'tblCellSpacing', 'jc', 'hidden',
        'ins', 'del', 'trPrChange',
    ],
}


def _reorder_el(parent):
    """Reordena hijos de elementos con orden estricto en el esquema OpenXML
    (rPr, tblPr, tcPr, trPr) según ECMA-376 y elimina duplicados del mismo tag."""
    if parent is None:
        return
    children = list(parent)
    if not children:
        return
    parent_tag = parent.tag.split('}')[-1] if '}' in parent.tag else parent.tag
    order = _SCHEMA_ORDERS.get(parent_tag)

    seen = {}
    for elem in children:
        tag_local = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
        if tag_local not in seen:
            seen[tag_local] = elem
        else:
            parent.remove(elem)
    remaining = list(parent)

    def key_fn(elem):
        tag_local = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
        if order is None:
            return 0
        return order.index(tag_local) if tag_local in order else 999

    sorted_children = sorted(remaining, key=key_fn)
    if sorted_children != remaining:
        for elem in remaining:
            parent.remove(elem)
        for elem in sorted_children:
            parent.append(elem)


def normalize_document_xml(doc):
    """Normalize all w:pPr, w:rPr, w:tblPr, w:tcPr, w:trPr elements
    across the document body to strictly satisfy OpenXML schema."""
    if doc is None or not hasattr(doc, 'element'):
        return
    wns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    for pPr in doc.element.body.xpath('.//w:pPr'):
        reorder_pPr(pPr)
    for rPr in doc.element.body.xpath('.//w:rPr'):
        _reorder_el(rPr)
    for tblPr in doc.element.body.xpath('.//w:tblPr'):
        _reorder_el(tblPr)
    for tcPr in doc.element.body.xpath('.//w:tcPr'):
        _reorder_el(tcPr)
    for trPr in doc.element.body.xpath('.//w:trPr'):
        _reorder_el(trPr)


def add_dynamic_page_number_to_footer(paragraph, doc=None):
    font_name = "Century Gothic"

    def add_field(p, field_text, bold=True):
        r_pr_inner = f'<w:rFonts {nsdecls("w")} w:ascii="{font_name}" w:hAnsi="{font_name}" w:cs="{font_name}"/><w:sz {nsdecls("w")} w:val="22"/><w:szCs {nsdecls("w")} w:val="22"/>'
        if bold:
            r_pr_inner += '<w:b/><w:bCs/>'
        r_pr = f'<w:rPr>{r_pr_inner}</w:rPr>'

        fld_begin = parse_xml(f'<w:r {nsdecls("w")}>{r_pr}<w:fldChar w:fldCharType="begin"/></w:r>')
        fld_instr = parse_xml(f'<w:r {nsdecls("w")}>{r_pr}<w:instrText xml:space="preserve"> {field_text} </w:instrText></w:r>')
        fld_sep   = parse_xml(f'<w:r {nsdecls("w")}>{r_pr}<w:fldChar w:fldCharType="separate"/></w:r>')
        fld_end   = parse_xml(f'<w:r {nsdecls("w")}>{r_pr}<w:fldChar w:fldCharType="end"/></w:r>')

        p._element.append(fld_begin)
        p._element.append(fld_instr)
        p._element.append(fld_sep)
        p._element.append(fld_end)

    r1 = paragraph.add_run("Página ")
    r1.font.name = font_name
    r1.font.size = Pt(11)

    add_field(paragraph, "PAGE", bold=True)

    r2 = paragraph.add_run(" de ")
    r2.font.name = font_name
    r2.font.size = Pt(11)

    add_field(paragraph, "NUMPAGES", bold=True)


def setup_footer_page_number(footer, doc=None):
    for p in footer.paragraphs: p.text = ""
    for table in footer.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs: p.text = ""
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_dynamic_page_number_to_footer(fp, doc=doc)


def force_single_column(section):
    sectPr = section._sectPr
    cols = sectPr.xpath('w:cols')
    if cols:
        cols[0].set(qn('w:num'), '1')
        for attr in ['w:space', 'w:equalWidth']:
            if qn(attr) in cols[0].attrib:
                del cols[0].attrib[qn(attr)]
    else:
        cols_xml = parse_xml(f'<w:cols {nsdecls("w")} w:num="1"/>')
        sectPr.append(cols_xml)


def apply_page_setup(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Cm(1)
    section.bottom_margin = Cm(1)
    section.left_margin = Cm(1)
    section.right_margin = Cm(1)
    section.header_distance = Cm(0.3)
    force_single_column(section)
    # Force vertical alignment to top
    sectPr = section._sectPr
    vAlign = sectPr.find(qn('w:vAlign'))
    if vAlign is not None:
        sectPr.remove(vAlign)


def get_all_paragraphs(doc):
    from docx.text.paragraph import Paragraph
    return [Paragraph(p, doc) for p in doc.element.body.xpath('.//w:p')]


def _prepend_run(paragraph, text, bold=False, font_name="Century Gothic", size_pt=11):
    """Add a run at the BEGINNING of the paragraph formatted to Century Gothic 11pt."""
    text = sanitize_xml_text(text)
    escaped = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
    new_r = parse_xml(f'<w:r {nsdecls("w")}><w:t xml:space="preserve">{escaped}</w:t></w:r>')
    first = paragraph._element.find(qn('w:r'))
    if first is not None:
        first.addprevious(new_r)
    else:
        paragraph._element.append(new_r)
    from docx.text.run import Run
    r = Run(new_r, paragraph)
    if bold:
        r.bold = True
    r.font.name = font_name
    r.font.size = Pt(size_pt)
    return r


def _remove_numpr(paragraph):
    pPr = paragraph._element.get_or_add_pPr()
    np = pPr.find(qn('w:numPr'))
    if np is not None:
        pPr.remove(np)


def _strip_prefix_from_runs(paragraph, length):
    wns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    for r_elem in paragraph._element.xpath('.//w:r'):
        if length <= 0:
            break
        t_elems = r_elem.findall(f'{{{wns}}}t')
        for t_elem in t_elems:
            if length <= 0:
                break
            t = t_elem.text or ""
            if not t:
                continue
            if len(t) <= length:
                length -= len(t)
                t_elem.text = ""
            else:
                t_elem.text = t[length:]
                length = 0


def _replace_paragraph_text_preserve_drawings(paragraph, new_text):
    wns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    t_elems = paragraph._element.xpath('.//w:t')
    if not t_elems:
        return
    t_elems[0].text = new_text
    for t_elem in t_elems[1:]:
        t_elem.text = ""


def apply_renumbering_and_ranges(doc, start_number):
    all_paras = get_all_paragraphs(doc)
    mapping = {}
    q_paras = []
    cur = start_number
    for p in all_paras:
        text = p.text.strip()
        if not text:
            continue
        m = re.match(r'^(\s*[\(\[\{]?)(\d+)(?:\s*[\.\)\]\}\-\:\/]+\s*|\s+)(?![\d\.])', text)
        if m:
            orig = int(m.group(2))
            if orig > 200:
                continue
            mapping[orig] = cur
            q_paras.append((p, orig, cur))
            cur += 1

    already_correct = q_paras and all(orig == new for _, orig, new in q_paras)

    if not already_correct:
        for p, orig_val, new_val in q_paras:
            mp = re.match(r'^(\s*[\(\[\{]?\d+(?:\s*[\.\)\]\}\-\:\/]+\s*|\s+))(?![\d\.])', p.text)
            if mp:
                _strip_prefix_from_runs(p, len(mp.group(1)))
                _prepend_run(p, f"{new_val}. ", bold=True)
                _remove_numpr(p)
    else:
        for p, orig_val, new_val in q_paras:
            text = p.text
            mp = re.match(r'^(\s*[\(\[\{]?\d+(?:\s*[\.\)\]\}\-\:\/]+\s*|\s+))(?![\d\.])', text)
            if mp:
                first = p.runs[0] if p.runs else None
                if first is None or not first.bold:
                    _strip_prefix_from_runs(p, len(mp.group(1)))
                    _prepend_run(p, f"{new_val}. ", bold=True)
                _remove_numpr(p)

    def replace_refs(t):
        if not mapping:
            return t
        t = re.sub(
            r'(?i)(de\s+la[s]?\s+(?:pregunta[s]?\s+)?)(\d+)(\s+(?:a|hasta)\s+la[s]?\s+(?:pregunta[s]?\s+)?)(\d+)',
            lambda m: m.group(1) + str(mapping.get(int(m.group(2)), int(m.group(2)))) + m.group(3) + str(mapping.get(int(m.group(4)), int(m.group(4)))), t)
        t = re.sub(
            r'(?i)(del\s+(?:pregunta[s]?\s+)?)(\d+)(\s+(?:al|hasta\s+el)\s+(?:pregunta[s]?\s+)?)(\d+)',
            lambda m: m.group(1) + str(mapping.get(int(m.group(2)), int(m.group(2)))) + m.group(3) + str(mapping.get(int(m.group(4)), int(m.group(4)))), t)
        t = re.sub(
            r'(?i)(pregunta[s]?\s+)(\d+)(\s+(?:a|al|hasta|y)\s+(?:pregunta[s]?\s+)?)(\d+)',
            lambda m: m.group(1) + str(mapping.get(int(m.group(2)), int(m.group(2)))) + m.group(3) + str(mapping.get(int(m.group(4)), int(m.group(4)))), t)
        return t
    for p in all_paras:
        old = p.text
        if not old.strip():
            continue
        new = replace_refs(old)
        if old != new:
            _replace_paragraph_text_preserve_drawings(p, new)
    return cur


def _resolve_autonumbering(doc):
    wns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    try:
        numbering_part = doc.part.numbering_part
        if numbering_part is None:
            return
        num_xml = numbering_part.element
    except:
        return

    abstract_levels = {}
    for abstract_num in num_xml.findall(f'{{{wns}}}abstractNum'):
        aid = int(abstract_num.get(qn('w:abstractNumId')))
        levels = {}
        for lvl in abstract_num.findall(f'{{{wns}}}lvl'):
            ilvl = int(lvl.get(qn('w:ilvl')) or '0')
            nf = lvl.find(f'{{{wns}}}numFmt')
            fmt = nf.get(qn('w:val')) if nf is not None else 'decimal'
            lt = lvl.find(f'{{{wns}}}lvlText')
            tpl = lt.get(qn('w:val')) if lt is not None else '%1.'
            st = lvl.find(f'{{{wns}}}start')
            start = int(st.get(qn('w:val'))) if st is not None else 1
            levels[ilvl] = (fmt, tpl, start)
        abstract_levels[aid] = levels

    num_to_abstract = {}
    for num_elem in num_xml.findall(f'{{{wns}}}num'):
        nid = int(num_elem.get(qn('w:numId')))
        aid_elem = num_elem.find(f'{{{wns}}}abstractNumId')
        if aid_elem is not None:
            num_to_abstract[nid] = int(aid_elem.get(qn('w:val')))

    counters = {}

    def level_value(num_id, ilvl, levels):
        k = (num_id, ilvl)
        if k not in counters:
            info = levels.get(ilvl)
            return info[2] if info else 1
        return counters[k]

    def fmt_num(val, fmt_name):
        if fmt_name in ('decimal', 'ordinal'):
            return str(val)
        elif fmt_name in ('upperLetter', 'upperAlpha'):
            return chr(ord('A') + val - 1) if 1 <= val <= 26 else str(val)
        elif fmt_name in ('lowerLetter', 'lowerAlpha'):
            return chr(ord('a') + val - 1) if 1 <= val <= 26 else str(val)
        return str(val)

    for p in get_all_paragraphs(doc):
        pPr = p._element.get_or_add_pPr()
        np = pPr.find(qn('w:numPr'))
        if np is None:
            continue

        nid_elem = np.find(qn('w:numId'))
        ilvl_elem = np.find(qn('w:ilvl'))
        if nid_elem is None:
            continue

        num_id = int(nid_elem.get(qn('w:val')))
        ilvl = int(ilvl_elem.get(qn('w:val'))) if ilvl_elem is not None else 0

        text = p.text.strip()
        if re.match(r'^(\d+|[a-eA-E])\s*[\.\)\]\}\-\:\/]', text):
            continue

        abstract_id = num_to_abstract.get(num_id)
        if abstract_id is None:
            continue
        levels = abstract_levels.get(abstract_id, {})
        level_info = levels.get(ilvl)
        if level_info is None:
            continue

        fmt, tpl, start_val = level_info

        cur_key = (num_id, ilvl)
        if cur_key not in counters:
            counters[cur_key] = start_val
        else:
            counters[cur_key] += 1

        for sub_ilvl in range(ilvl + 1, 10):
            sub_key = (num_id, sub_ilvl)
            if sub_key in counters:
                info = levels.get(sub_ilvl)
                counters[sub_key] = info[2] if info else 1

        def resolve_ph(m):
            idx = int(m.group(1))
            target_ilvl = idx - 1
            val = level_value(num_id, target_ilvl, levels)
            info = levels.get(target_ilvl)
            lvl_fmt = info[0] if info else 'decimal'
            return fmt_num(val, lvl_fmt)

        prefix = re.sub(r'%(\d+)', resolve_ph, tpl) + ' '
        _prepend_run(p, prefix, bold=True)
        _remove_numpr(p)


def remove_indents(paragraph):
    pPr = paragraph._element.get_or_add_pPr()
    ind = pPr.find(qn('w:ind'))
    if ind is not None:
        pPr.remove(ind)
    paragraph.paragraph_format.left_indent = Pt(0)
    paragraph.paragraph_format.right_indent = Pt(0)
    paragraph.paragraph_format.first_line_indent = Pt(0)


def _clear_paragraph_style_level(paragraph):
    pPr = paragraph._element.get_or_add_pPr()
    style_elem = pPr.find(qn('w:pStyle'))
    if style_elem is not None:
        pPr.remove(style_elem)
    outline_elem = pPr.find(qn('w:outlineLvl'))
    if outline_elem is not None:
        pPr.remove(outline_elem)


def _normalize_case(text):
    text = text.strip()
    if not text:
        return ""
    if text.isupper() and len(text) > 2:
        connectors = {'y', 'e', 'o', 'u', 'de', 'del', 'la', 'las', 'los', 'el', 'en', 'con', 'por', 'para', 'a', 'al', 'of', 'and', 'or', 'in', 'on', 'at', 'to', 'for', 'with'}
        parts = re.split(r'(\s+|[/,\-])', text.lower())
        res = []
        is_first = True
        for p in parts:
            if not p:
                continue
            if re.match(r'^\s+|[/,\-]$', p):
                res.append(p)
            else:
                if is_first:
                    res.append(p.capitalize())
                    is_first = False
                elif p in connectors:
                    res.append(p.lower())
                else:
                    res.append(p.capitalize())
        return "".join(res)
    return text


def _add_styled_run(paragraph, text, bold=False, size_pt=11, font_name="Century Gothic"):
    if not text:
        return None
    r = paragraph.add_run(text)
    r.bold = bold
    r.font.name = font_name
    r.font.size = Pt(size_pt)
    return r


def _remove_tabs_from_pPr(p):
    """Elimina paradas de tabulación XML (<w:tabs>) y tabuladores del párrafo."""
    wns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    pPr = p._element.get_or_add_pPr()
    tabs = pPr.find(f'{{{wns}}}tabs')
    if tabs is not None:
        pPr.remove(tabs)
    for t_elem in list(p._element.xpath('.//w:tab')):
        parent = t_elem.getparent()
        if parent is not None:
            parent.remove(t_elem)


def _ensure_ends_with_period(paragraph):
    text = paragraph.text.strip()
    if text and not text.endswith(('.', ':', ';', '!', '?', ')', ']', '}')):
        wns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        t_elems = paragraph._element.xpath('.//w:t')
        if t_elems:
            for t_elem in reversed(t_elems):
                if t_elem.text and t_elem.text.strip():
                    t_elem.text = t_elem.text.rstrip() + '.'
                    break
            else:
                t_elems[-1].text = (t_elems[-1].text or "") + "."
        else:
            paragraph.add_run('.')


BULLET_CHARS_RE = re.compile(r'^\s*[\u2022\u25CF\u25A0\u25A1\u25AA\u25AB\uF0B7\u2013\u2014\u00B7\u25A2\u25C6\u25C7\u25E6\u2023\u2043•□▪·–\*]\s*')

def is_bullet_paragraph(paragraph):
    pPr = paragraph._element.find(qn('w:pPr'))
    if pPr is not None and pPr.find(qn('w:numPr')) is not None:
        return True
    text = paragraph.text.strip()
    if BULLET_CHARS_RE.match(text):
        return True
    return False


def _format_competencia_or_componente_paragraph(p, force_lang=None, font_name="Century Gothic"):
    raw_text = p.text.replace('\r', ' ').replace('\n', ' ').replace('\t', ' ').strip()
    text = re.sub(r'\s+', ' ', raw_text)
    if not text or _has_drawing(p._element):
        return False

    # Check if paragraph contains BOTH Competencia/Competence AND Componente/Component on the same line/paragraph
    m_comp = re.search(r'\b(Competencia|Competence)\s*[:\-]?\s*(.*?)(?=\b(?:Componente|Component)\b|$)', text, re.IGNORECASE)
    m_compo = re.search(r'\b(Componente|Component)\s*[:\-]?\s*(.*?)$', text, re.IGNORECASE)

    if m_comp and m_compo and m_comp.start() < m_compo.start() and m_compo.start() > 0:
        c1_label, c1_val = m_comp.group(1), m_comp.group(2).strip()
        c2_label, c2_val = m_compo.group(1), m_compo.group(2).strip()

        p.text = ""
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        lang1 = force_lang or ("en" if 'competence' in c1_label.lower() else "es")
        lbl1 = "Competence" if lang1 == "en" else "Competencia"
        val1 = _normalize_case(c1_val)
        if val1 and not val1.endswith(('.', ':', ';', '!', '?')):
            val1 += '.'
        _add_styled_run(p, f"{lbl1}: ", bold=True, size_pt=11, font_name=font_name)
        _add_styled_run(p, val1, bold=False, size_pt=11, font_name=font_name)

        p_elem = p._element
        parent = p_elem.getparent()
        if parent is not None and c2_val.strip(' .:-'):
            p2_xml = parse_xml(f'<w:p {nsdecls("w")}/>')
            p_elem.addnext(p2_xml)
            from docx.text.paragraph import Paragraph
            p2_obj = Paragraph(p2_xml, p._parent)
            lang2 = force_lang or ("en" if 'component' in c2_label.lower() else "es")
            lbl2 = "Component" if lang2 == "en" else "Componente"
            val2 = _normalize_case(c2_val)
            if val2 and not val2.endswith(('.', ':', ';', '!', '?')):
                val2 += '.'
            _add_styled_run(p2_obj, f"{lbl2}: ", bold=True, size_pt=11, font_name=font_name)
            _add_styled_run(p2_obj, val2, bold=False, size_pt=11, font_name=font_name)
            p2_obj.alignment = WD_ALIGN_PARAGRAPH.LEFT
            remove_indents(p2_obj)
            set_single_line_spacing(p2_obj)
        return True

    m = re.match(r'^(Competencia|Competence|Componente|Component|Habilidades|Habilidad|Nivel|Level|Desempeño|Aprendizaje|Afirmación|Estándar)\s*[:\-]?\s*(.*)$', text, re.IGNORECASE)
    if m:
        raw_val = m.group(2).strip()
        if raw_val.lower().strip(' .:-') in ['es', 'en', '']:
            raw_val = ""
        raw_label = m.group(1).lower()
        if force_lang == 'en':
            label = "Competence" if ('compet' in raw_label and 'competencia' not in raw_label) or 'competence' in raw_label else ("Component" if 'component' in raw_label else raw_label.capitalize())
        elif force_lang == 'es':
            label = "Competencia" if ('compet' in raw_label and 'competence' not in raw_label) or 'competencia' in raw_label else ("Componente" if 'componente' in raw_label or 'component' in raw_label else raw_label.capitalize())
        else:
            if 'competence' in raw_label:
                label = "Competence"
            elif 'componente' in raw_label:
                label = "Componente"
            elif 'competencia' in raw_label:
                label = "Competencia"
            elif 'component' in raw_label:
                label = "Component"
            elif 'habilidad' in raw_label:
                label = "Habilidad" if raw_label == 'habilidad' else "Habilidades"
            elif 'nivel' in raw_label:
                label = "Nivel"
            elif 'level' in raw_label:
                label = "Level"
            else:
                label = raw_label.capitalize()

        _remove_tabs_from_pPr(p)
        remove_indents(p)
        set_single_line_spacing(p)
        p.text = ""
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if raw_val and raw_val.strip(' .:-') != '':
            content = _normalize_case(raw_val)
            if content and not content.endswith(('.', ':', ';', '!', '?')):
                content += '.'
            _add_styled_run(p, f"{label}: ", bold=True, size_pt=11, font_name=font_name)
            _add_styled_run(p, content, bold=False, size_pt=11, font_name=font_name)
        else:
            _add_styled_run(p, f"{label}:", bold=True, size_pt=11, font_name=font_name)
        return True
    return False


def format_paragraph(paragraph, doc_ref):
    text = paragraph.text.strip()
    if not text:
        return

    is_bullet = is_bullet_paragraph(paragraph)
    if not is_bullet:
        # Strip all indents by default for clean flush-left alignment of normal paragraphs
        remove_indents(paragraph)
    else:
        # For bullet paragraphs, ensure alignment is LEFT and a clean hanging indent exists
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pPr = paragraph._element.get_or_add_pPr()
        ind = pPr.find(qn('w:ind'))
        if ind is None:
            ind_elem = parse_xml(f'<w:ind {nsdecls("w")} w:left="360" w:hanging="240"/>')
            pPr.append(ind_elem)

    # 1. Clean up redundant input document headers and teacher metadata
    is_redundant_header = (
        re.search(r'(NOMBRE|APELLIDO|ESTUDIANTE|ALUMNO|CURSO|GRADO|FECHA|CÓDIGO|CODIGO|NAME|DATE|CODE)\s*:\s*[_]{2,}', text, re.IGNORECASE) or
        re.search(r'^\s*(EVALUACI[OÓ]N\s+(BIMESTRAL|DIAGN[OÓ]STICA|FINAL|PARCIAL|DE\s+SUFICIENCIA|ACAD[EÉ]MICA|SUMATIVA)|EXAMEN\s+DE|CUESTIONARIO\s+DE|PRUEBA\s+DE)', text, re.IGNORECASE) or
        re.search(r'^\s*(GRADO|CURSO)\s+(SEXTO|S[EÉ]PTIMO|OCTAVO|NOVENO|D[EÉ]CIMO|ONCE|PRIMERO|SEGUNDO|TERCERO|CUARTO|QUINTO|TRANSICI[OÓ]N|JARD[IÍ]N|PREJARD[IÍ]N|\d+[\u00b0°]?)', text, re.IGNORECASE) or
        re.search(r'^\s*(DOCENTE|PROFESOR|PROFESORA|ASIGNATURA|MATERIA)\s*:\s*', text, re.IGNORECASE) or
        re.search(r'^\s*[\(\[\{]?\s*(EJE\s+TEM[AÁ]TICO|EJE|TEM[AÁ]TICA|TEMA|INDICADOR(\s+DE\s+(LOGRO|DESEMPE[NÑ]O))?|EST[AÁ]NDAR|DESEMPE[NÑ]O|AFIRMACI[OÓ]N|EVIDENCIA|RESPUESTA\s+CORRECTA|CLAVE(\s+DE\s+RESPUESTA)?)\s*[:\-]\s*', text, re.IGNORECASE)
    )

    if is_redundant_header:
        # Don't delete if it's Competencia/Competence/Componente/Component/Habilidad/Nivel or a Question or Bullet!
        if not re.match(r'^(Competencia|Competence|Componente|Component|Habilidades|Habilidad|Nivel|Level|Desempeño|Aprendizaje|Estándar|Eje)\s*[:\-]?\s*', text, re.IGNORECASE) and not re.match(r'^\s*\d+[\.\)]', text) and not is_bullet:
            if _has_drawing(paragraph._element):
                for t in paragraph._element.xpath('.//w:t'):
                    t.text = ""
            else:
                paragraph.text = ""
            return

    if _format_competencia_or_componente_paragraph(paragraph):
        return
    # Option letter A–E — format prefix (supports A), a), A., a., (A), A)text)
    mo = re.match(r'^(\s*[\(\[\{]?([a-eA-E])\s*[\.\)\]\}\-\:\/]\s*)(?!\d)', text)
    if mo:
        runs = paragraph.runs
        first = runs[0] if runs else None
        already_good = (
            first is not None
            and first.text.strip().upper() == (mo.group(2).upper() + ".")
            and first.bold
        )
        if not already_good:
            _strip_prefix_from_runs(paragraph, len(mo.group(1)))
            _prepend_run(paragraph, f"{mo.group(2).upper()}. ", bold=True)
            _remove_numpr(paragraph)
        remove_indents(paragraph)
        _ensure_ends_with_period(paragraph)
        return
    # Question number — only bold if not already
    mn = re.match(r'^(\s*[\(\[\{]?(\d+)(?:\s*[\.\)\]\}\-\:\/]+\s*|\s+))(?![\d\.])', text)
    if mn:
        runs = paragraph.runs
        first = runs[0] if runs else None
        already_bold = first is not None and first.bold
        if not already_bold:
            _strip_prefix_from_runs(paragraph, len(mn.group(1)))
            _prepend_run(paragraph, f"{mn.group(2)}. ", bold=True)
        _remove_numpr(paragraph)
        remove_indents(paragraph)


def reset_letter_sequence(doc):
    paras = get_all_paragraphs(doc)
    letter_index = 0

    for p in paras:
        text = p.text.strip()
        if not text:
            continue

        mn = re.match(r'^(\s*[\(\[\{]?)(\d+)(?:\s*[\.\)\]\}\-\:\/]+\s*|\s+)(?![\d\.])', text)
        if mn and int(mn.group(2)) <= 200:
            letter_index = 0
            continue

        mo = re.match(r'^(\s*[\(\[\{]?)([a-eA-E])(\s*[\.\)\]\}\-\:\/]\s*)', text)
        if mo:
            current = mo.group(2).upper()
            expected = chr(ord('A') + letter_index) if letter_index < 5 else None

            if expected and current != expected:
                full_prefix = mo.group(0)
                _strip_prefix_from_runs(p, len(full_prefix))
                _prepend_run(p, f"{expected}. ", bold=True)
                _remove_numpr(p)

            letter_index += 1
            continue


def set_single_line_spacing(paragraph):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0
    pPr = paragraph._element.get_or_add_pPr()

    # Remove contextualSpacing if present so Word doesn't alter spacing unpredictably
    contextual = pPr.find(qn('w:contextualSpacing'))
    if contextual is not None:
        pPr.remove(contextual)

    sp = pPr.find(qn('w:spacing'))
    if sp is None:
        sp = parse_xml(f'<w:spacing {nsdecls("w")} w:line="240" w:lineRule="auto" w:before="0" w:after="0"/>')
        pPr.append(sp)
    else:
        sp.set(qn('w:line'), '240')
        sp.set(qn('w:lineRule'), 'auto')
        sp.set(qn('w:before'), '0')
        sp.set(qn('w:after'), '0')
        for key in ['beforeAutospacing', 'afterAutospacing', 'beforeLines', 'afterLines']:
            if qn(f'w:{key}') in sp.attrib:
                del sp.attrib[qn(f'w:{key}')]

    reorder_pPr(pPr)


def _has_drawing(para_element):
    """Check if a paragraph element contains any drawing, picture, shape or embedded object."""
    wns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    vns = 'urn:schemas-microsoft-com:vml'
    return (
        len(para_element.findall(f'.//{{{wns}}}drawing')) > 0 or
        len(para_element.findall(f'.//{{{wns}}}pict')) > 0 or
        len(para_element.findall(f'.//{{{vns}}}shape')) > 0 or
        len(para_element.findall(f'.//{{{vns}}}imagedata')) > 0 or
        len(para_element.findall(f'.//{{{wns}}}object')) > 0
    )


def inject_list_definitions(doc, start_number=1):
    wns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    try:
        num_part = doc.part.numbering_part
    except NotImplementedError:
        return None
    
    if num_part is None:
        return None

    num_xml = num_part.element
    if not num_xml.xpath('w:abstractNum[@w:abstractNumId="9000"]'):
        # Decimal list abstractNum (Questions) - strict OpenXML child order inside w:lvl
        abs_dec = parse_xml(f'''
            <w:abstractNum {nsdecls('w')} w:abstractNumId="9000">
                <w:multiLevelType w:val="hybridMultilevel"/>
                <w:lvl w:ilvl="0">
                    <w:start w:val="1"/>
                    <w:numFmt w:val="decimal"/>
                    <w:suff w:val="space"/>
                    <w:lvlText w:val="%1."/>
                    <w:lvlJc w:val="left"/>
                    <w:pPr>
                        <w:spacing w:before="0" w:after="0"/>
                        <w:ind w:left="360" w:hanging="360"/>
                    </w:pPr>
                    <w:rPr>
                        <w:b w:val="1"/>
                        <w:rFonts w:ascii="Century Gothic" w:hAnsi="Century Gothic" w:cs="Century Gothic"/>
                        <w:sz w:val="22"/>
                        <w:szCs w:val="22"/>
                    </w:rPr>
                </w:lvl>
            </w:abstractNum>
        ''')
        
        # UpperLetter list abstractNum (Options - same left indent as Questions)
        abs_alpha = parse_xml(f'''
            <w:abstractNum {nsdecls('w')} w:abstractNumId="9001">
                <w:multiLevelType w:val="hybridMultilevel"/>
                <w:lvl w:ilvl="0">
                    <w:start w:val="1"/>
                    <w:numFmt w:val="upperLetter"/>
                    <w:suff w:val="space"/>
                    <w:lvlText w:val="%1."/>
                    <w:lvlJc w:val="left"/>
                    <w:pPr>
                        <w:spacing w:before="0" w:after="0"/>
                        <w:ind w:left="360" w:hanging="360"/>
                    </w:pPr>
                    <w:rPr>
                        <w:b w:val="1"/>
                        <w:rFonts w:ascii="Century Gothic" w:hAnsi="Century Gothic" w:cs="Century Gothic"/>
                        <w:sz w:val="22"/>
                        <w:szCs w:val="22"/>
                    </w:rPr>
                </w:lvl>
            </w:abstractNum>
        ''')
        
        # Bullet list abstractNum (Bullets)
        abs_bullet = parse_xml(f'''
            <w:abstractNum {nsdecls('w')} w:abstractNumId="9002">
                <w:multiLevelType w:val="hybridMultilevel"/>
                <w:lvl w:ilvl="0">
                    <w:start w:val="1"/>
                    <w:numFmt w:val="bullet"/>
                    <w:suff w:val="space"/>
                    <w:lvlText w:val="•"/>
                    <w:lvlJc w:val="left"/>
                    <w:pPr>
                        <w:spacing w:before="0" w:after="0"/>
                        <w:ind w:left="360" w:hanging="240"/>
                    </w:pPr>
                    <w:rPr>
                        <w:rFonts w:ascii="Century Gothic" w:hAnsi="Century Gothic" w:cs="Century Gothic"/>
                        <w:sz w:val="22"/>
                        <w:szCs w:val="22"/>
                    </w:rPr>
                </w:lvl>
            </w:abstractNum>
        ''')
        
        first_num = num_xml.find(qn('w:num'))
        if first_num is not None:
            first_num.addprevious(abs_dec)
            first_num.addprevious(abs_alpha)
            first_num.addprevious(abs_bullet)
        else:
            num_xml.append(abs_dec)
            num_xml.append(abs_alpha)
            num_xml.append(abs_bullet)
            
        num_dec = parse_xml(f'''
            <w:num {nsdecls('w')} w:numId="9000">
                <w:abstractNumId w:val="9000"/>
                <w:lvlOverride w:ilvl="0">
                    <w:startOverride w:val="{start_number}"/>
                </w:lvlOverride>
            </w:num>
        ''')
        num_xml.append(num_dec)

        # Pre-create numId=9100 as the first option-list container so that
        # any option paragraph referencing 9100 always finds a valid definition.
        # Subsequent questions will create 9101, 9102, ... as needed.
        num_opt_base = parse_xml(f'''
            <w:num {nsdecls('w')} w:numId="9100">
                <w:abstractNumId w:val="9001"/>
                <w:lvlOverride w:ilvl="0">
                    <w:startOverride w:val="1"/>
                </w:lvlOverride>
            </w:num>
        ''')
        num_xml.append(num_opt_base)
        
        # Pre-create numId=9002 for bullets
        num_bullet = parse_xml(f'''
            <w:num {nsdecls('w')} w:numId="9002">
                <w:abstractNumId w:val="9002"/>
            </w:num>
        ''')
        num_xml.append(num_bullet)
    else:
        # Update startOverride on existing num 9000
        existing_num = num_xml.xpath('w:num[@w:numId="9000"]')
        if existing_num:
            n_elem = existing_num[0]
            lvl_ovr = n_elem.xpath('w:lvlOverride')
            if not lvl_ovr:
                lvl_ovr_elem = parse_xml(f'<w:lvlOverride {nsdecls("w")} w:ilvl="0"><w:startOverride w:val="{start_number}"/></w:lvlOverride>')
                n_elem.append(lvl_ovr_elem)
            else:
                st = lvl_ovr[0].xpath('w:startOverride')
                if st:
                    st[0].set(qn('w:val'), str(start_number))
                else:
                    st_elem = parse_xml(f'<w:startOverride {nsdecls("w")} w:val="{start_number}"/>')
                    lvl_ovr[0].append(st_elem)

        # Ensure numId=9100 (base options container) always exists
        existing_9100 = num_xml.xpath('w:num[@w:numId="9100"]')
        if not existing_9100:
            num_opt_base = parse_xml(f'''
                <w:num {nsdecls('w')} w:numId="9100">
                    <w:abstractNumId w:val="9001"/>
                    <w:lvlOverride w:ilvl="0">
                        <w:startOverride w:val="1"/>
                    </w:lvlOverride>
                </w:num>
            ''')
            num_xml.append(num_opt_base)
            
        # Ensure numId=9002 (bullet list container) always exists
        existing_9002 = num_xml.xpath('w:num[@w:numId="9002"]')
        if not existing_9002:
            num_bullet = parse_xml(f'''
                <w:num {nsdecls('w')} w:numId="9002">
                    <w:abstractNumId w:val="9002"/>
                </w:num>
            ''')
            num_xml.append(num_bullet)

    return True


def _strip_header_footer_parts(docx_path):
    """Elimina del ZIP todos los archivos de header/footer y sus relaciones,
    para evitar que Word COM detecte corrupción en sub-documentos procesados."""
    if not os.path.exists(docx_path):
        return
    from lxml import etree
    import zipfile, io
    rns = 'http://schemas.openxmlformats.org/package/2006/relationships'

    # Leer el ZIP original
    with zipfile.ZipFile(docx_path, 'r') as zin:
        all_names = set(zin.namelist())
        data = {name: zin.read(name) for name in all_names}

    # Identificar archivos de header/footer
    header_files = [n for n in all_names if ('header' in n.lower() or 'footer' in n.lower()) and n.endswith('.xml') and 'rels' not in n]
    header_rels_files = [n for n in all_names if ('header' in n.lower() or 'footer' in n.lower()) and n.endswith('.rels')]

    if not header_files and not header_rels_files:
        return

    # Modificar document.xml.rels para eliminar relaciones de header/footer
    doc_rels_path = 'word/_rels/document.xml.rels'
    if doc_rels_path in data:
        rels_xml = etree.fromstring(data[doc_rels_path])
        for rel in list(rels_xml.findall(f'{{{rns}}}Relationship')):
            rtype = rel.get('Type', '')
            target = rel.get('Target', '')
            if 'header' in rtype.split('/')[-1] or 'footer' in rtype.split('/')[-1] or 'header' in target or 'footer' in target:
                rels_xml.remove(rel)
        data[doc_rels_path] = etree.tostring(rels_xml, xml_declaration=True, encoding='UTF-8', standalone=True)

    # Limpiar [Content_Types].xml de entradas de header/footer
    ct_path = '[Content_Types].xml'
    if ct_path in data:
        ct_xml = etree.fromstring(data[ct_path])
        ct_ns = 'http://schemas.openxmlformats.org/package/2006/content-types'
        for ov in list(ct_xml.findall(f'{{{ct_ns}}}Override')):
            pn = ov.get('PartName', '')
            for hf in header_files:
                if hf in pn or ('/' + hf) in pn:
                    ct_xml.remove(ov)
                    break
        data[ct_path] = etree.tostring(ct_xml, xml_declaration=True, encoding='UTF-8', standalone=True)

    # Eliminar header/footer XML y sus .rels del paquete
    files_to_remove = set(header_files + header_rels_files)
    for f in files_to_remove:
        data.pop(f, None)

    # Reconstruir el ZIP
    with zipfile.ZipFile(docx_path, 'w', zipfile.ZIP_DEFLATED) as zout:
        for name in sorted(data.keys()):
            zout.writestr(name, data[name])


def _rebuild_zip(docx_path):
    """Reconstruye el ZIP para eliminar corrupción de compresión/estructura.
    [Content_Types].xml debe ir primero y sin comprimir según OPC."""
    if not os.path.exists(docx_path):
        return
    tmp_path = docx_path + '.rebuild.zip'
    try:
        with zipfile.ZipFile(docx_path, 'r') as zin:
            names = sorted(zin.namelist())
            with zipfile.ZipFile(tmp_path, 'w', zipfile.ZIP_DEFLATED) as zout:
                for name in names:
                    data = zin.read(name)
                    if name == '[Content_Types].xml':
                        zout.writestr(zipfile.ZipInfo(name), data, zipfile.ZIP_STORED)
                    else:
                        zout.writestr(name, data, zipfile.ZIP_DEFLATED)
        shutil.move(tmp_path, docx_path)
    except Exception:
        if os.path.exists(tmp_path):
            try: os.remove(tmp_path)
            except: pass


def _clean_rsid_attributes(docx_path):
    """Elimina atributos w:rsidR/w:rsidP/w:rsidRPr que python-docx agrega
    durante save() y que Word a veces rechaza."""
    if not os.path.exists(docx_path):
        return
    from lxml import etree
    import zipfile as zf
    tmp_path = docx_path + '.rsid.zip'
    wns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    rsid_attrs = [
        f'{{{wns}}}rsidR', f'{{{wns}}}rsidRPr', f'{{{wns}}}rsidP',
        f'{{{wns}}}rsidRDefault', f'{{{wns}}}rsidSect',
    ]
    try:
        with zf.ZipFile(docx_path, 'r') as zin:
            data = {name: zin.read(name) for name in zin.namelist()}
        for xml_name in list(data.keys()):
            if xml_name.startswith('word/') and not xml_name.startswith('word/_rels/') and xml_name.endswith('.xml'):
                try:
                    doc_xml = etree.fromstring(data[xml_name])
                    for elem in doc_xml.iter():
                        for attr in rsid_attrs:
                            if attr in elem.attrib:
                                del elem.attrib[attr]
                    for rsids in doc_xml.findall(f'{{{wns}}}rsids'):
                        rsids.getparent().remove(rsids)
                    data[xml_name] = etree.tostring(doc_xml, xml_declaration=True,
                                                     encoding='UTF-8', standalone=True)
                except Exception:
                    pass
        with zf.ZipFile(tmp_path, 'w', zf.ZIP_DEFLATED) as zout:
            for name in sorted(data.keys()):
                d = data[name]
                if name == '[Content_Types].xml':
                    zout.writestr(zf.ZipInfo(name), d, zf.ZIP_STORED)
                else:
                    zout.writestr(name, d, zf.ZIP_DEFLATED)
        shutil.move(tmp_path, docx_path)
    except Exception:
        if os.path.exists(tmp_path):
            try: os.remove(tmp_path)
            except: pass


def _clean_orphaned_header_footer_rels(docx_path):
    """Elimina del document.xml.rels y del ZIP los headers/footers
    que no tienen referencia activa en los sectPr del document.xml."""
    if not os.path.exists(docx_path):
        return
    from lxml import etree
    import zipfile as zf
    tmp_path = docx_path + '.clean.zip'
    wns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    try:
        with zf.ZipFile(docx_path, 'r') as zin:
            names = set(zin.namelist())
            data = {name: zin.read(name) for name in names}

        # Obtener los IDs de header/footer referenciados en los sectPr
        doc_xml = etree.fromstring(data['word/document.xml'])
        active_hf_refs = set()
        for sectPr in doc_xml.xpath('.//w:sectPr', namespaces={'w': wns}):
            for ref in sectPr.findall(f'{{{wns}}}headerReference'):
                rid = ref.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                if rid:
                    active_hf_refs.add(rid)
            for ref in sectPr.findall(f'{{{wns}}}footerReference'):
                rid = ref.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                if rid:
                    active_hf_refs.add(rid)

        # Identificar relaciones de header/footer huérfanas en document.xml.rels
        doc_rels_path = 'word/_rels/document.xml.rels'
        rns = 'http://schemas.openxmlformats.org/package/2006/relationships'
        rels_to_remove = set()

        if doc_rels_path in data:
            rels_xml = etree.fromstring(data[doc_rels_path])
            for rel in list(rels_xml.findall(f'{{{rns}}}Relationship')):
                rtype = rel.get('Type', '')
                rid = rel.get('Id', '')
                target = rel.get('Target', '')
                short_type = rtype.split('/')[-1] if '/' in rtype else rtype
                if short_type in ('header', 'footer') and rid not in active_hf_refs:
                    rels_to_remove.add(rid)
                    target_path = 'word/' + target.lstrip('/')
                    if target_path in data:
                        data.pop(target_path, None)
                    rels_xml.remove(rel)
            data[doc_rels_path] = etree.tostring(rels_xml, xml_declaration=True,
                                                  encoding='UTF-8', standalone=True)

        # Limpiar [Content_Types].xml SOLO de los headers/footers eliminados
        ct_path = '[Content_Types].xml'
        if ct_path in data:
            ct_xml = etree.fromstring(data[ct_path])
            ct_ns = 'http://schemas.openxmlformats.org/package/2006/content-types'
            removed_filenames = set()
            for ov in list(ct_xml.findall(f'{{{ct_ns}}}Override')):
                pn = ov.get('PartName', '').lstrip('/')
                if ('/header' in pn or '/footer' in pn) and pn not in data:
                    ct_xml.remove(ov)
            data[ct_path] = etree.tostring(ct_xml, xml_declaration=True,
                                            encoding='UTF-8', standalone=True)

        # Reconstruir ZIP limpio
        with zf.ZipFile(tmp_path, 'w', zf.ZIP_DEFLATED) as zout:
            for name in sorted(data.keys()):
                d = data[name]
                if name == '[Content_Types].xml':
                    zout.writestr(zf.ZipInfo(name), d, zf.ZIP_STORED)
                else:
                    zout.writestr(name, d, zf.ZIP_DEFLATED)
        shutil.move(tmp_path, docx_path)
    except Exception:
        if os.path.exists(tmp_path):
            try: os.remove(tmp_path)
            except: pass


def _ensure_sectPr_is_last(doc):
    """Mueve w:sectPr al final del body si no lo está ya (requisito OpenXML)."""
    if doc is None or not hasattr(doc, 'element'):
        return
    wns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    body = doc.element.body
    body_sectPr = body.find(f'{{{wns}}}sectPr')
    if body_sectPr is not None and list(body)[-1] != body_sectPr:
        body.remove(body_sectPr)
        body.append(body_sectPr)


def strip_leading_tabs(para):
    wns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    for t_elem in para._element.xpath('.//w:t'):
        if not t_elem.text:
            continue
        cleaned = t_elem.text.lstrip('\t')
        if cleaned != t_elem.text:
            t_elem.text = cleaned
        if t_elem.text:
            break


def apply_native_lists_to_final_doc(final_doc, start_offset=0):
    start_num = start_offset + 1
    if inject_list_definitions(final_doc, start_number=start_num) is None:
        return
    try:
        num_xml = final_doc.part.numbering_part.element
    except (NotImplementedError, AttributeError):
        return
    if num_xml is None:
        return
    q_num_id = 9000
    o_num_id_base = 9100
    o_num_id = o_num_id_base
        
    paras = get_all_paragraphs(final_doc)
    for p in paras:
        strip_leading_tabs(p)
        set_single_line_spacing(p)
        text = p.text.strip()
        
        # Check Option
        mo = re.match(r'^(\s*[\(\[\{]?([a-eA-E])\s*[\.\)\]\}\-\:\/]\s)(?!\d)', text)
        if mo:
            _strip_prefix_from_runs(p, len(mo.group(1)))
            pPr = p._element.get_or_add_pPr()
            numPr = pPr.find(qn('w:numPr'))
            if numPr is not None:
                pPr.remove(numPr)
            ind = pPr.find(qn('w:ind'))
            if ind is not None:
                pPr.remove(ind)
            numPr = parse_xml(f'<w:numPr {nsdecls("w")}><w:ilvl w:val="0"/><w:numId w:val="{o_num_id}"/></w:numPr>')
            pPr.append(numPr)
            ind_elem = parse_xml(f'<w:ind {nsdecls("w")} w:left="360" w:hanging="360"/>')
            pPr.append(ind_elem)
            reorder_pPr(pPr)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            strip_leading_tabs(p)
            _ensure_ends_with_period(p)
            continue
            
        # Check Question
        mn = re.match(r'^(\s*[\(\[\{]?(\d+)(?:\s*[\.\)\]\}\-\:\/]+\s*|\s+))(?![\d\.])', text)
        if mn:
            _strip_prefix_from_runs(p, len(mn.group(1)))
            pPr = p._element.get_or_add_pPr()
            numPr = pPr.find(qn('w:numPr'))
            if numPr is not None:
                pPr.remove(numPr)
            ind = pPr.find(qn('w:ind'))
            if ind is not None:
                pPr.remove(ind)
            numPr = parse_xml(f'<w:numPr {nsdecls("w")}><w:ilvl w:val="0"/><w:numId w:val="{q_num_id}"/></w:numPr>')
            pPr.append(numPr)
            ind_elem = parse_xml(f'<w:ind {nsdecls("w")} w:left="360" w:hanging="360"/>')
            pPr.append(ind_elem)
            reorder_pPr(pPr)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            strip_leading_tabs(p)
            
            # Restart options for this question
            if num_xml is not None:
                o_num_id += 1
                new_num = parse_xml(f'''
                    <w:num {nsdecls('w')} w:numId="{o_num_id}">
                        <w:abstractNumId w:val="9001"/>
                        <w:lvlOverride w:ilvl="0">
                            <w:startOverride w:val="1"/>
                        </w:lvlOverride>
                    </w:num>
                ''')
                num_xml.append(new_num)
            continue

        # Check Bullet
        mb = BULLET_CHARS_RE.match(p.text)
        if mb:
            _strip_prefix_from_runs(p, len(mb.group(0)))
            pPr = p._element.get_or_add_pPr()
            numPr = pPr.find(qn('w:numPr'))
            if numPr is not None:
                pPr.remove(numPr)
            ind = pPr.find(qn('w:ind'))
            if ind is not None:
                pPr.remove(ind)
            numPr = parse_xml(f'<w:numPr {nsdecls("w")}><w:ilvl w:val="0"/><w:numId w:val="9002"/></w:numPr>')
            pPr.append(numPr)
            ind_elem = parse_xml(f'<w:ind {nsdecls("w")} w:left="360" w:hanging="240"/>')
            pPr.append(ind_elem)
            reorder_pPr(pPr)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            strip_leading_tabs(p)
            _ensure_ends_with_period(p)
            continue


def _fix_numbering_level_fonts(doc, font_name="Century Gothic"):
    """Fix fonts in numbering.xml levels safely without breaking bullet/symbol levels."""
    try:
        if hasattr(doc.part, 'numbering_part') and doc.part.numbering_part is not None:
            wns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
            num_xml = doc.part.numbering_part.element
            for lvl in num_xml.xpath('.//w:lvl'):
                numFmt = lvl.find(f'{{{wns}}}numFmt')
                fmt_val = (numFmt.get(qn('w:val')) if numFmt is not None else '').lower()
                lvlText = lvl.find(f'{{{wns}}}lvlText')
                t_val = lvlText.get(qn('w:val')) if lvlText is not None else ''

                if fmt_val in ('bullet', 'none') or any(ord(c) >= 0xF000 or c in '•·–—□▪' for c in t_val):
                    continue

                rPr = lvl.find(f'{{{wns}}}rPr')
                if rPr is None:
                    continue

                rFonts = rPr.find(f'{{{wns}}}rFonts')
                if rFonts is not None:
                    ascii_f = (rFonts.get(qn('w:ascii')) or '').lower()
                    if any(s in ascii_f for s in ['symbol', 'wingdings', 'webdings', 'marlett', 'dingbats']):
                        continue
                    rFonts.set(qn('w:ascii'), font_name)
                    rFonts.set(qn('w:hAnsi'), font_name)
                    rFonts.set(qn('w:cs'), font_name)
                    for theme_attr in ['w:asciiTheme', 'w:hAnsiTheme', 'w:eastAsiaTheme', 'w:csTheme', 'w:theme']:
                        if qn(theme_attr) in rFonts.attrib:
                            del rFonts.attrib[qn(theme_attr)]
    except Exception:
        pass


def apply_formatting_to_document(doc):
    font_name = "Century Gothic"
    from docx.text.paragraph import Paragraph
    
    # Process all paragraphs in the document body
    for p_elem in doc.element.body.xpath('.//w:p'):
        para = Paragraph(p_elem, doc)
        text = para.text.strip()
        
        is_comp = bool(re.match(r'^(Competencia|Competence|Componente|Component)\s*[:\-]', text, re.IGNORECASE))
        
        if not _has_drawing(p_elem):
            if is_comp:
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
        format_paragraph(para, doc)
        set_single_line_spacing(para)
        strip_leading_tabs(para)
        
        if is_comp:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Empty separator lines: force 2pt (w:val="4") for paper saving per AGENTS.md rule
        if not para.text.strip() and not _has_drawing(p_elem):
            wns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
            for r_elem in p_elem.xpath('.//w:r'):
                rPr = r_elem.find(f'{{{wns}}}rPr')
                if rPr is None:
                    rPr = parse_xml(f'<w:rPr {nsdecls("w")}/>')
                    r_elem.insert(0, rPr)
                sz = rPr.find(f'{{{wns}}}sz')
                if sz is None:
                    sz = parse_xml(f'<w:sz {nsdecls("w")} w:val="4"/>')
                    rPr.append(sz)
                else:
                    sz.set(qn('w:val'), '4')
                szCs = rPr.find(f'{{{wns}}}szCs')
                if szCs is None:
                    szCs = parse_xml(f'<w:szCs {nsdecls("w")} w:val="4"/>')
                    rPr.append(szCs)
                else:
                    szCs.set(qn('w:val'), '4')

    # Format all runs in the document BODY XML (excluding headers/footers)
    wns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    for r_elem in doc.element.body.xpath('.//w:r'):
        rPr = r_elem.find(f'{{{wns}}}rPr')
        if rPr is None:
            rPr = parse_xml(f'<w:rPr {nsdecls("w")}/>')
            r_elem.insert(0, rPr)
        
        rFonts = rPr.find(f'{{{wns}}}rFonts')
        if rFonts is None:
            rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{font_name}" w:hAnsi="{font_name}" w:cs="{font_name}"/>')
            rPr.append(rFonts)
        else:
            ascii_f = (rFonts.get(qn('w:ascii')) or '').lower()
            if not any(sym in ascii_f for sym in ['symbol', 'wingdings', 'webdings', 'marlett']):
                rFonts.set(qn('w:ascii'), font_name)
                rFonts.set(qn('w:hAnsi'), font_name)
                rFonts.set(qn('w:cs'), font_name)
                for theme_attr in ['w:asciiTheme', 'w:hAnsiTheme', 'w:eastAsiaTheme', 'w:csTheme', 'w:theme']:
                    if qn(theme_attr) in rFonts.attrib:
                        del rFonts.attrib[qn(theme_attr)]

        sz = rPr.find(f'{{{wns}}}sz')
        if sz is None:
            sz = parse_xml(f'<w:sz {nsdecls("w")} w:val="22"/>')
            rPr.append(sz)
        else:
            sz.set(qn('w:val'), '22')

        szCs = rPr.find(f'{{{wns}}}szCs')
        if szCs is None:
            szCs = parse_xml(f'<w:szCs {nsdecls("w")} w:val="22"/>')
            rPr.append(szCs)
        else:
            szCs.set(qn('w:val'), '22')

    _fix_numbering_level_fonts(doc, font_name)


def _safe_remove_para(p):
    """Elimina el párrafo. Retorna True si se eliminó, False si se conservó
    (por ejemplo cuando es el último bloque de una celda de tabla)."""
    try:
        p_elem = p._element
        parent = p_elem.getparent()
        if parent is not None:
            # No eliminar el último bloque de una celda de tabla: w:tc exige
            # al menos un elemento de bloque (w:p/w:tbl). Dejarlo vacío produce
            # XML inválido que hace fallar a Word ("contenido ilegible").
            anc = parent
            while anc is not None:
                if anc.tag.endswith('}tc'):
                    block_children = [c for c in anc if c.tag.endswith('}p') or c.tag.endswith('}tbl')]
                    if p_elem in block_children and len(block_children) <= 1:
                        return False
                    break
                anc = anc.getparent()
            parent.remove(p_elem)
            return True
    except Exception:
        pass
    return False


def remove_blank_lines_between_question_parts(doc):
    all_paras = get_all_paragraphs(doc)
    i = 0
    while i < len(all_paras):
        p = all_paras[i]
        text = p.text.strip()
        
        is_question = bool(re.match(r'^\s*\d+[\.\)]', text))
        is_option = bool(re.match(r'^\s*([a-eA-E])[\.\)]', text))
        is_comp = bool(re.match(r'^(Competencia|Competence|Componente|Component)\s*:', text, re.IGNORECASE))
        
        # If this is a question, option, or competencia/componente, remove following blank lines (unless they contain images/drawings)
        if is_question or is_option or is_comp:
            j = i + 1
            while j < len(all_paras):
                np = all_paras[j]
                if not np.text.strip():
                    if not _has_drawing(np._element):
                        _safe_remove_para(np)
                        j += 1
                    else:
                        # Paragraph contains an image/drawing -> keep intact and stop removing!
                        break
                else:
                    # Non-empty paragraph reached
                    break
        i += 1


def replace_xml_text(elements_list, replacements):
    wns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    dns = 'http://schemas.openxmlformats.org/drawingml/2006/main'
    for el in elements_list:
        for t in el.iter(f'{{{wns}}}t'):
            if t.text:
                for k, v in replacements.items():
                    if k in t.text:
                        t.text = sanitize_xml_text(t.text.replace(k, str(v)))
        for t in el.iter(f'{{{dns}}}t'):
            if t.text:
                for k, v in replacements.items():
                    if k in t.text:
                        t.text = sanitize_xml_text(t.text.replace(k, str(v)))


def count_questions_in_doc(doc):
    c = 0
    for para in doc.paragraphs:
        if re.match(r'^(\d+)(?:\s*[\.\)\]\}\-\:\/]+\s*|\s+)(?![\d\.])', para.text.strip()):
            c += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if re.match(r'^(\d+)(?:\s*[\.\)\]\}\-\:\/]+\s*|\s+)(?![\d\.])', para.text.strip()):
                        c += 1
    return c


def make_header_tables_inline(header):
    for table in header.tables:
        tblPr = table._element.xpath('w:tblPr')
        if tblPr:
            tblpPr = tblPr[0].xpath('w:tblpPr')
            if tblpPr:
                tblPr[0].remove(tblpPr[0])


def clear_header_completely(header):
    header.is_linked_to_previous = False
    for p in list(header.paragraphs):
        p.text = ""
    for t in list(header.tables):
        try:
            t._element.getparent().remove(t._element)
        except:
            pass
    header._element.clear()


def apply_section0_page_setup(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)
    force_single_column(section)
    # Force vertical alignment to top
    sectPr = section._sectPr
    vAlign = sectPr.find(qn('w:vAlign'))
    if vAlign is not None:
        sectPr.remove(vAlign)


def apply_subsequent_page_setup(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)
    section.header_distance = Cm(0.5)
    section.footer_distance = Cm(0.8)
    force_single_column(section)
    # Force vertical alignment to top
    sectPr = section._sectPr
    vAlign = sectPr.find(qn('w:vAlign'))
    if vAlign is not None:
        sectPr.remove(vAlign)


def replace_in_all(tpl, replacements_map):
    """Replace placeholders in the entire template (body + headers)."""
    # Replace in body (user already moved header content here)
    replace_xml_text([tpl.element.body], replacements_map)
    # Also replace in any leftover headers (defensive)
    for s in tpl.sections:
        for h in [s.header, s.first_page_header]:
            replace_xml_text(list(h._element), replacements_map)


def convert_competencia_tables_to_paragraphs(doc):
    """
    Si una tabla contiene Competencia / Competence / Componente / Component (ej: celdas o layout de tabla),
    extrae el contenido a párrafos normales alineados a la izquierda y elimina la tabla.
    """
    from docx.text.paragraph import Paragraph
    for tbl in list(doc.tables):
        tbl_text = ""
        for row in tbl.rows:
            for cell in row.cells:
                tbl_text += " " + cell.text.strip()
        tbl_text = tbl_text.strip()
        
        if re.search(r'\b(Competencia|Competence|Componente|Component)\s*[:\-]', tbl_text, re.IGNORECASE):
            tbl_elem = tbl._element
            parent = tbl_elem.getparent()
            if parent is not None:
                for row in tbl.rows:
                    row_txt = " ".join([c.text.strip() for c in row.cells if c.text.strip()])
                    if row_txt:
                        new_p = parse_xml(f'<w:p {nsdecls("w")}/>')
                        tbl_elem.addprevious(new_p)
                        p_obj = Paragraph(new_p, doc)
                        p_obj.text = row_txt
                        _format_competencia_or_componente_paragraph(p_obj)
                parent.remove(tbl_elem)


def split_inline_competencia_and_componente(doc):
    """
    Si en un solo párrafo vienen pegados Competencia/Competence y Componente/Component
    (ej: 'Competence: Linguistic/sociolinguistic COMPONENT: Comprehension'),
    se dividen en 2 párrafos separados en líneas independientes.
    """
    from docx.text.paragraph import Paragraph
    for p in get_all_paragraphs(doc):
        text = p.text.replace('\t', ' ').strip()
        m_comp = re.search(r'\b(Competencia|Competence)\s*[:\-]?\s*(.*?)(?=\b(?:Componente|Component)\b|$)', text, re.IGNORECASE)
        m_compo = re.search(r'\b(Componente|Component)\s*[:\-]?\s*(.*?)$', text, re.IGNORECASE)
        
        # Si el párrafo contiene AMBOS marcadores a la vez
        if m_comp and m_compo and m_comp.start() < m_compo.start():
            comp_label = m_comp.group(1)
            comp_val = m_comp.group(2).strip()
            
            compo_label = m_compo.group(1)
            compo_val = m_compo.group(2).strip()
            
            p_elem = p._element
            parent = p_elem.getparent()
            if parent is not None:
                if comp_val.strip(' .:-'):
                    p1_xml = parse_xml(f'<w:p {nsdecls("w")}/>')
                    p_elem.addprevious(p1_xml)
                    p1_obj = Paragraph(p1_xml, doc)
                    p1_obj.text = f"{comp_label}: {comp_val}"
                    _format_competencia_or_componente_paragraph(p1_obj)
                
                if compo_val.strip(' .:-'):
                    p2_xml = parse_xml(f'<w:p {nsdecls("w")}/>')
                    p_elem.addprevious(p2_xml)
                    p2_obj = Paragraph(p2_xml, doc)
                    p2_obj.text = f"{compo_label}: {compo_val}"
                    _format_competencia_or_componente_paragraph(p2_obj)
                
                parent.remove(p_elem)


def reorder_competencia_before_question(doc):
    """
    Garantiza que Competencia, Componente, Habilidad y Nivel aparezcan SIEMPRE
    ANTES del enunciado de la pregunta (ej. '54. Cuando acoplamos...'), nunca debajo del enunciado ni encima de A-D.
    """
    all_paras = get_all_paragraphs(doc)
    i = 0
    while i < len(all_paras):
        p = all_paras[i]
        text = p.text.strip()

        # Verificar si p es un enunciado de pregunta (ej. "54.")
        m_q = re.match(r'^\s*(\d+)[\.\)]', text)
        if m_q:
            comp_paras_to_move = []
            j = i + 1
            while j < len(all_paras):
                np = all_paras[j]
                np_text = np.text.strip()

                # Detener al encontrar opción A-E o la siguiente pregunta
                if re.match(r'^\s*[\(\[\{]?([a-eA-E])\s*[\.\)\]\}\-\:\/]', np_text) or re.match(r'^\s*\d+[\.\)]', np_text):
                    break

                # Detectar Competencia, Componente, Habilidad, Nivel, etc.
                if re.match(r'^(Competencia|Competence|Componente|Component|Habilidades|Habilidad|Nivel|Level|Desempeño|Aprendizaje|Afirmación|Estándar)\b', np_text, re.IGNORECASE):
                    comp_paras_to_move.append(np)

                j += 1

            if comp_paras_to_move:
                q_elem = p._element
                for c_p in comp_paras_to_move:
                    q_elem.addprevious(c_p._element)
                all_paras = get_all_paragraphs(doc)

        i += 1


def _format_bullet_item_clean(p):
    text = p.text.strip()
    if not text or _has_drawing(p._element):
        return

    m_bul = BULLET_CHARS_RE.match(text)
    if m_bul:
        bul_sym = m_bul.group(0).strip()
        body_text = text[m_bul.end():].strip()

        if body_text:
            c_text = _normalize_case(body_text)
            if c_text and not c_text.endswith(('.', ':', ';', '!', '?', ')', ']')):
                c_text += '.'

            rPr_xml = None
            runs = p.runs
            if runs:
                rPr_elem = runs[0]._element.find(qn('w:rPr'))
                if rPr_elem is not None:
                    rPr_xml = copy.deepcopy(rPr_elem)

            p.text = ""
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_single_line_spacing(p)

            pPr = p._element.get_or_add_pPr()
            ind = pPr.find(qn('w:ind'))
            if ind is not None:
                pPr.remove(ind)
            ind_elem = parse_xml(f'<w:ind {nsdecls("w")} w:left="360" w:hanging="240"/>')
            pPr.append(ind_elem)

            r1 = p.add_run(f"{bul_sym}  ")
            r1.bold = False
            r1.font.size = Pt(11)
            if rPr_xml is not None:
                rFonts = rPr_xml.find(qn('w:rFonts'))
                if rFonts is not None:
                    a_font = rFonts.get(qn('w:ascii')) or ''
                    if any(s in a_font.lower() for s in ['symbol', 'wingdings', 'webdings', 'marlett']):
                        r1.font.name = a_font

            if c_text:
                r2 = p.add_run(c_text)
                r2.bold = False
                r2.font.name = "Century Gothic"
                r2.font.size = Pt(11)


def process_habilidades_and_bullets(doc):
    """
    1. Fusiona el desglose de habilidades compartidas en el bloque en una sola
       línea (sin viñeta) con el formato exacto:
       "Habilidades: Comparar, Definir, Relacionar."
       (habilidades separadas por comas y terminadas en punto).
    2. Formatea las viñetas independientes con inicial en mayúscula
       (Sentence Case, ej: Comparar., Definir.).
    3. Garantiza espacio limpio y libre de colisiones entre el icono de la viñeta y el texto.
    """
    all_paras = get_all_paragraphs(doc)
    i = 0
    while i < len(all_paras):
        p = all_paras[i]
        text = p.text.strip()

        m_hab = re.match(r'^(Habilidades|Habilidad)\s*[:\-]?\s*(.*)$', text, re.IGNORECASE)
        if m_hab:
            val = m_hab.group(2).strip()
            if val.lower().strip(' .:-') in ['es', 'en', '']:
                val = ""

            bullet_items = []
            j = i + 1
            while j < len(all_paras):
                np = all_paras[j]
                np_text = np.text.strip()
                if is_bullet_paragraph(np):
                    bullet_items.append(np)
                elif not np_text and not _has_drawing(np._element):
                    pass
                else:
                    break
                j += 1

            # Recolectar las habilidades (valor inline + viñetas) para escribirlas
            # en una sola línea separadas por comas.
            skill_texts = []
            if val:
                skill_texts.append(_normalize_case(val).rstrip('.'))
            for b_p in bullet_items:
                bt = b_p.text.strip()
                m_bul = BULLET_CHARS_RE.match(bt)
                if m_bul:
                    bt = bt[m_bul.end():].strip()
                bt = _normalize_case(bt).rstrip('.')
                if bt:
                    skill_texts.append(bt)

            if skill_texts:
                # Formato exacto (sin viñeta): Habilidades: Comparar, Definir, Relacionar.
                correct_label = "Habilidades" if len(skill_texts) >= 2 else "Habilidad"
                body_text = ', '.join(skill_texts)
                if not body_text.endswith(('.', ':', ';', '!', '?')):
                    body_text += '.'

                p.text = ""
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                remove_indents(p)
                set_single_line_spacing(p)
                _add_styled_run(p, f"{correct_label}: ", bold=True, size_pt=11, font_name="Century Gothic")
                _add_styled_run(p, body_text, bold=False, size_pt=11, font_name="Century Gothic")

                for b_p in bullet_items:
                    _safe_remove_para(b_p)
                all_paras = get_all_paragraphs(doc)
            else:
                # Sin habilidades enumeradas: solo la etiqueta
                correct_label = "Habilidades" if len(bullet_items) >= 2 else "Habilidad"
                p.text = ""
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                remove_indents(p)
                set_single_line_spacing(p)
                _add_styled_run(p, f"{correct_label}:", bold=True, size_pt=11, font_name="Century Gothic")
                for b_p in bullet_items:
                    _format_bullet_item_clean(b_p)

        elif is_bullet_paragraph(p):
            _format_bullet_item_clean(p)

        i += 1


def process_competencias_and_componentes(doc):
    # Always use Century Gothic — the post-merge pass will enforce it anyway
    font_name = "Century Gothic"

    # Step 0: Extraer Competencia/Componente atrapados en tablas a párrafos normales libres
    convert_competencia_tables_to_paragraphs(doc)

    # Step 0.1: Dividir si vienen pegados en la misma línea
    split_inline_competencia_and_componente(doc)

    all_paras = get_all_paragraphs(doc)

    # Step 1: Normalize all Competencia / Competence / Componente / Component paragraphs
    for p in all_paras:
        _format_competencia_or_componente_paragraph(p, font_name=font_name)

    # Step 2: Ensure order (Competencia/Competence first, Componente/Component second) and Language consistency
    all_paras = get_all_paragraphs(doc)
    i = 0
    while i < len(all_paras) - 1:
        p1 = all_paras[i]
        p2 = all_paras[i + 1]
        t1 = p1.text.replace('\t', ' ').strip()
        t2 = p2.text.replace('\t', ' ').strip()
        
        m1 = re.match(r'^(Competencia|Competence|Componente|Component)\s*[:\-]?\s*(.*)$', t1, re.IGNORECASE)
        m2 = re.match(r'^(Competencia|Competence|Componente|Component)\s*[:\-]?\s*(.*)$', t2, re.IGNORECASE)
        
        if m1 and m2:
            l1 = m1.group(1).lower()
            l2 = m2.group(1).lower()
            
            # Detect language of the pair: if any is explicitly Spanish ('competencia'/'componente'), enforce Spanish.
            # If both are English ('competence'/'component'), enforce English.
            if 'competencia' in l1 or 'componente' in l1 or 'competencia' in l2 or 'componente' in l2:
                pair_lang = 'es'
            elif 'competence' in l1 or 'component' in l1 or 'competence' in l2 or 'component' in l2:
                pair_lang = 'en'
            else:
                pair_lang = 'es'
                
            is_comp1 = ('componente' in l1 or 'component' in l1)
            is_comp2 = ('competencia' in l2 or 'competence' in l2)
            
            if is_comp1 and is_comp2:
                # Swap paragraphs so Competencia/Competence comes first
                p1_text = p1.text
                p1.text = p2.text
                p2.text = p1_text
                
            _format_competencia_or_componente_paragraph(p1, force_lang=pair_lang, font_name=font_name)
            _format_competencia_or_componente_paragraph(p2, force_lang=pair_lang, font_name=font_name)
        elif m1:
            l1 = m1.group(1).lower()
            single_lang = 'en' if ('competence' in l1 or l1 == 'component') else 'es'
            _format_competencia_or_componente_paragraph(p1, force_lang=single_lang, font_name=font_name)
            
        i += 1

    # Step 3: Remove blank lines directly under Competencia/Competence or Componente/Component
    all_paras = get_all_paragraphs(doc)
    for i, p in enumerate(all_paras):
        t_str = p.text.strip()
        if re.match(r'^(Competencia|Competence|Componente|Component)\s*:', t_str, re.IGNORECASE):
            j = i + 1
            while j < len(all_paras):
                np = all_paras[j]
                if not np.text.strip() and not _has_drawing(np._element):
                    _safe_remove_para(np)
                    j += 1
                else:
                    break

    # Step 4: Block-level Deduplication (Only remove if the PAIR is identical across consecutive questions)
    all_paras = get_all_paragraphs(doc)
    last_pair = (None, None)
    
    i = 0
    while i < len(all_paras):
        p = all_paras[i]
        t_str = p.text.strip()
        
        m_comp = re.match(r'^(Competencia|Competence)\s*:\s*(.*)$', t_str, re.IGNORECASE)
        if m_comp:
            comp_val = m_comp.group(2).strip().lower()
            compo_val = None
            compo_p = None
            
            if i + 1 < len(all_paras):
                m_co = re.match(r'^(Componente|Component)\s*:\s*(.*)$', all_paras[i + 1].text.strip(), re.IGNORECASE)
                if m_co:
                    compo_val = m_co.group(2).strip().lower()
                    compo_p = all_paras[i + 1]
        
            current_pair = (comp_val, compo_val)
            if current_pair == last_pair and comp_val is not None:
                # Remove duplicate pair!
                if not _has_drawing(p._element):
                    _safe_remove_para(p)
                if compo_p is not None and not _has_drawing(compo_p._element):
                    _safe_remove_para(compo_p)
            else:
                last_pair = current_pair

        i += 1

    # Step 5: Process Habilidad / Habilidades singular/plural + bullet item casing & spacing
    process_habilidades_and_bullets(doc)


def ensure_proper_spacing_between_questions(doc):
    """
    1. Elimina líneas vacías superfluas entre PART X, Competencia/Componente y la pregunta.
    2. Garantiza exactamente 1 línea en blanco al terminar las opciones de una pregunta 
       antes de iniciar el siguiente bloque (PART X, Competencia o nueva Pregunta).
    """
    # 1. Eliminar líneas en blanco consecutivas (más de 1 en blanco seguidas)
    all_paras = get_all_paragraphs(doc)
    i = 0
    while i < len(all_paras) - 1:
        p1 = all_paras[i]
        p2 = all_paras[i + 1]
        if p1._element.getparent() == doc.element.body and p2._element.getparent() == doc.element.body:
            if not p1.text.strip() and not _has_drawing(p1._element) and not p2.text.strip() and not _has_drawing(p2._element):
                _safe_remove_para(p1)
                all_paras = get_all_paragraphs(doc)
                continue
        i += 1

    # 2. Eliminar cualquier línea en blanco entre PART X, Competencia, Componente y el enunciado de la pregunta
    all_paras = get_all_paragraphs(doc)
    i = 0
    while i < len(all_paras) - 1:
        p = all_paras[i]
        t_str = p.text.strip()
        is_header_elem = (
            bool(re.match(r'^(PART|PARTE)\s+\d+[\s\:\-]', t_str, re.IGNORECASE)) or
            bool(re.match(r'^(Competencia|Competence|Componente|Component)\s*:', t_str, re.IGNORECASE))
        )
        if is_header_elem:
            np = all_paras[i + 1]
            if not np.text.strip() and not _has_drawing(np._element):
                if _safe_remove_para(np):
                    all_paras = get_all_paragraphs(doc)
                    continue
        i += 1

    # 3. Eliminar cualquier línea en blanco entre las opciones (A, B, C, D) de las preguntas
    all_paras = get_all_paragraphs(doc)
    i = 0
    while i < len(all_paras) - 1:
        p = all_paras[i]
        text = p.text.strip()
        is_opt = bool(re.match(r'^\s*[\(\[\{]?([a-eA-E])\s*[\.\)\]\}\-\:\/]', text))
        if is_opt:
            j = i + 1
            while j < len(all_paras):
                np = all_paras[j]
                np_text = np.text.strip()
                if not np_text and not _has_drawing(np._element):
                    if _safe_remove_para(np):
                        all_paras = get_all_paragraphs(doc)
                    else:
                        break
                else:
                    break
        i += 1

    # 3.5 Eliminar línea en blanco entre el enunciado de la pregunta y la primera opción (A, B, C, D)
    all_paras = get_all_paragraphs(doc)
    i = 0
    while i < len(all_paras) - 1:
        p = all_paras[i]
        if bool(re.match(r'^\s*\d+[\.\)]', p.text.strip())):
            np = all_paras[i + 1]
            np_text = np.text.strip()
            if not np_text and not _has_drawing(np._element):
                nnp = all_paras[i + 2] if i + 2 < len(all_paras) else None
                if nnp and bool(re.match(r'^\s*[\(\[\{]?([a-eA-E])\s*[\.\)\]\}\-\:\/]', nnp.text.strip())):
                    if _safe_remove_para(np):
                        all_paras = get_all_paragraphs(doc)
                        continue
        i += 1

    # 4. Garantizar exactamente 1 línea en blanco al terminar una pregunta/opciones antes del nuevo bloque
    all_paras = get_all_paragraphs(doc)
    i = 0
    while i < len(all_paras):
        p = all_paras[i]
        text = p.text.strip()
        
        is_part_header = bool(re.match(r'^(PART|PARTE)\s+\d+[\s\:\-]', text, re.IGNORECASE))
        is_comp_header = bool(re.match(r'^(Competencia|Competence)\s*:', text, re.IGNORECASE))
        is_question_num = bool(re.match(r'^\s*\d+[\.\)]', text))
        
        prev_text = all_paras[i - 1].text.strip() if i > 0 else ""
        prev_is_comp_or_part = (
            bool(re.match(r'^(Competencia|Competence|Componente|Component)\s*:', prev_text, re.IGNORECASE)) or
            bool(re.match(r'^(PART|PARTE)\s+\d+[\s\:\-]', prev_text, re.IGNORECASE))
        )
        
        is_start_of_block = is_part_header or is_comp_header or (is_question_num and not prev_is_comp_or_part)
        
        if is_start_of_block and i > 0:
            if prev_text and not prev_is_comp_or_part:
                p._element.addprevious(parse_xml(f'<w:p {nsdecls("w")}/>'))
                all_paras = get_all_paragraphs(doc)
                i += 1
        i += 1


def _merge_docx_with_rels(master_doc, sub_doc, add_break=False):
    rel_map = {}
    master_pkg = master_doc.part.package
    existing_parts = {part.partname: part for part in master_pkg.iter_parts()}
    existing_part_names = set(existing_parts.keys())

    for rel_id, rel in list(sub_doc.part.rels.items()):
        if "image" in rel.target_ref:
            try:
                part = rel.target_part
                if part.partname in existing_part_names:
                    from docx.parts.image import ImagePart
                    ext = part.partname.ext
                    new_name = master_pkg.next_partname('/word/media/image%%d.%s' % ext)
                    new_part = ImagePart.from_image(part.image, new_name)
                    new_rel_id = master_doc.part.relate_to(new_part, rel.reltype)
                    existing_parts[new_name] = new_part
                    existing_part_names.add(new_name)
                else:
                    new_rel_id = master_doc.part.relate_to(part, rel.reltype)
                    existing_parts[part.partname] = part
                    existing_part_names.add(part.partname)
                rel_map[rel_id] = new_rel_id
            except Exception as e:
                print(f"[GESA] Error mapeando imagen {rel_id}: {e}")
        elif "hyperlink" in rel.target_ref:
            try:
                target_part = rel.target_part
                if target_part is not None and target_part.partname is not None:
                    new_rel_id = master_doc.part.relate_to(target_part, rel.reltype)
                    rel_map[rel_id] = new_rel_id
            except Exception:
                pass

    # ── Merge numbering definitions from sub_doc into master_doc ──────────
    # This prevents orphaned numId references that cause "contenido ilegible".
    _merge_numbering(master_doc, sub_doc)

    # CRITICAL: Insert elements BEFORE the body's sectPr, not after it.
    # OpenXML spec requires sectPr to be the LAST element of body.
    # body.append() would push elements after sectPr, corrupting the file.
    wns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    body = master_doc.element.body
    body_sectPr = body.find(f'{{{wns}}}sectPr')

    for element in sub_doc.element.body:
        if element.tag.endswith('sectPr'):
            continue
        new_elem = copy.deepcopy(element)
        if rel_map:
            unmapped_embeds = set()
            for r_elem in new_elem.xpath('.//*[@r:embed]'):
                old_embed = r_elem.get(qn('r:embed'))
                if old_embed in rel_map:
                    r_elem.set(qn('r:embed'), rel_map[old_embed])
                else:
                    unmapped_embeds.add(old_embed)
                    r_elem.attrib.pop(qn('r:embed'), None)
            if unmapped_embeds:
                for r_elem in list(new_elem.iter()):
                    if r_elem.tag.endswith('}r') or r_elem.tag == 'r':
                        drawing = r_elem.find(qn('w:drawing'))
                        if drawing is not None:
                            embeds_in_drawing = set()
                            for emb_elem in drawing.xpath('.//*[@r:embed]'):
                                embeds_in_drawing.add(emb_elem.get(qn('r:embed')))
                            if embeds_in_drawing and embeds_in_drawing.issubset(unmapped_embeds):
                                r_elem.getparent().remove(r_elem)
        if body_sectPr is not None:
            body_sectPr.addprevious(new_elem)
        else:
            body.append(new_elem)

    if add_break:
        # Also insert the page break before sectPr
        if body_sectPr is not None:
            br_p = parse_xml(f'<w:p {nsdecls("w")}><w:r><w:br w:type="page"/></w:r></w:p>')
            body_sectPr.addprevious(br_p)
        else:
            master_doc.add_page_break()


def _merge_numbering(master_doc, sub_doc):
    """Transfer all abstractNum and num definitions from sub_doc to master_doc,
    remapping IDs to avoid collisions. Updates numId references in sub_doc body
    before the elements are copied over."""
    wns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    try:
        sub_num_part = sub_doc.part.numbering_part
        if sub_num_part is None:
            return
        sub_num_xml = sub_num_part.element
    except (NotImplementedError, AttributeError):
        return

    try:
        master_num_part = master_doc.part.numbering_part
        if master_num_part is None:
            return
        master_num_xml = master_num_part.element
    except (NotImplementedError, AttributeError):
        return

    # Find highest existing IDs in master to avoid collision
    existing_abs_ids = set()
    for an in master_num_xml.findall(f'{{{wns}}}abstractNum'):
        aid = an.get(qn('w:abstractNumId'))
        if aid is not None:
            try: existing_abs_ids.add(int(aid))
            except: pass

    existing_num_ids = set()
    for n in master_num_xml.findall(f'{{{wns}}}num'):
        nid = n.get(qn('w:numId'))
        if nid is not None:
            try: existing_num_ids.add(int(nid))
            except: pass

    GESA_RESERVED_LOWER = 9000
    GESA_RESERVED_UPPER = 9999

    max_abs_id = max(existing_abs_ids, default=0)
    max_num_id = max(existing_num_ids, default=0)

    if max_abs_id >= GESA_RESERVED_LOWER and max_abs_id < GESA_RESERVED_UPPER:
        max_abs_id = GESA_RESERVED_UPPER
    if max_num_id >= GESA_RESERVED_LOWER and max_num_id < GESA_RESERVED_UPPER:
        max_num_id = GESA_RESERVED_UPPER

    # Build mapping: old sub_doc abstractNumId → new ID in master
    abs_id_map = {}
    for an in sub_num_xml.findall(f'{{{wns}}}abstractNum'):
        old_id = an.get(qn('w:abstractNumId'))
        if old_id is None:
            continue
        old_id_int = int(old_id)
        max_abs_id += 1
        new_id = max_abs_id
        abs_id_map[old_id_int] = new_id

        new_an = copy.deepcopy(an)
        new_an.set(qn('w:abstractNumId'), str(new_id))
        # Insert before the first w:num in master (abstractNums must precede nums)
        first_num = master_num_xml.find(f'{{{wns}}}num')
        if first_num is not None:
            first_num.addprevious(new_an)
        else:
            master_num_xml.append(new_an)

    # Build mapping: old sub_doc numId → new ID in master
    num_id_map = {}
    for n in sub_num_xml.findall(f'{{{wns}}}num'):
        old_nid = n.get(qn('w:numId'))
        if old_nid is None:
            continue
        old_nid_int = int(old_nid)
        max_num_id += 1
        new_nid = max_num_id
        num_id_map[old_nid_int] = new_nid

        new_n = copy.deepcopy(n)
        new_n.set(qn('w:numId'), str(new_nid))
        # Update abstractNumId reference inside the num element
        abs_ref = new_n.find(f'{{{wns}}}abstractNumId')
        if abs_ref is not None:
            old_abs_ref = abs_ref.get(qn('w:val'))
            if old_abs_ref is not None:
                mapped = abs_id_map.get(int(old_abs_ref))
                if mapped is not None:
                    abs_ref.set(qn('w:val'), str(mapped))
        master_num_xml.append(new_n)

    # Now update all numId references in the sub_doc body XML
    if num_id_map:
        for numId_elem in sub_doc.element.body.xpath('.//w:numId'):
            old_val = numId_elem.get(qn('w:val'))
            if old_val is not None:
                try:
                    old_int = int(old_val)
                    if old_int in num_id_map:
                        numId_elem.set(qn('w:val'), str(num_id_map[old_int]))
                except (ValueError, TypeError):
                    pass



def strip_leading_empty_paras_and_breaks(doc):
    body = doc.element.body
    for elem in list(body):
        if elem.tag.endswith('}p') or elem.tag == 'p':
            for br in elem.xpath('.//w:br[@w:type="page"]'):
                parent = br.getparent()
                if parent is not None:
                    parent.remove(br)
            text = ''.join(elem.xpath('.//w:t/text()')).strip()
            if not text and not _has_drawing(elem):
                body.remove(elem)
            else:
                break


def strip_trailing_empty_paras_and_breaks(doc):
    body = doc.element.body
    for elem in reversed(list(body)):
        if elem.tag.endswith('}p') or elem.tag == 'p':
            for br in elem.xpath('.//w:br[@w:type="page"]'):
                parent = br.getparent()
                if parent is not None:
                    parent.remove(br)
            text = ''.join(elem.xpath('.//w:t/text()')).strip()
            if not text and not _has_drawing(elem):
                body.remove(elem)
            else:
                break


def strip_section_breaks(doc):
    """Remove ALL inline sectPr elements (section breaks inside paragraph pPr or body).
    This merges all sections into one so content flows continuously without
    forced page or section breaks from the source documents."""
    wns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    body = doc.element.body
    body_sectPr = body.find(f'{{{wns}}}sectPr')
    for sectPr in list(body.xpath('.//w:sectPr')):
        if sectPr == body_sectPr:
            continue
        parent = sectPr.getparent()
        if parent is not None and not parent.tag.endswith('body'):
            parent.remove(sectPr)


# ── MAIN MERGE ───────────────────────────────────────────────

def merge_docx_with_guaranteed_header(template_path, file_list, output_path, config_data, start_offset=0, stop_check=None):
    try:
        return _merge_impl(template_path, file_list, output_path, config_data, start_offset, stop_check)
    except Exception:
        import traceback
        tb = traceback.format_exc()
        raise RuntimeError(
            f"Error en merge_docx_with_guaranteed_header:\n{tb}"
        )


def _merge_impl(template_path, file_list, output_path, config_data, start_offset=0, stop_check=None):
    grade_top = config_data.get('grade', config_data.get('grade_clean', ''))
    p_c_val = config_data.get('p_c_value', grade_top)
    replacements_map = {
        "(EDU_LEVEL)": str(config_data['level']).upper(),
        "(GRADE)": str(grade_top).upper(),
        "(TERM)": str(config_data['period']).upper(),
        "(SESSION)": str(config_data['session_code']).upper(),
        "(DATE)": str(config_data['date']).upper(),
        "(P_C)": str(p_c_val),
    }

    num_files = len([f for f in file_list if os.path.exists(f)])
    eval_prefix = "Evaluación" if num_files <= 1 else "Evaluaciones"

    title_context = {k: config_data.get(k, '') for k in ['grade_clean', 'period', 'session_code', 'year', 'level']}
    title_context['grade'] = config_data.get('grade', title_context['grade_clean'])
    title_context['session'] = title_context['session_code']

    title_template = config_data.get('title_template', '')
    if not title_template:
        title_template = f"{eval_prefix} de Suficiencia Académica - {{grade_clean}} - {{period}} - {{session_code}} - {{year}}"
    else:
        if title_template.startswith("Evaluación de") or title_template.startswith("Evaluaciones de"):
            title_template = re.sub(r'^Evaluaci\u00f3n(es)?\s+de', f'{eval_prefix} de', title_template)

    expanded_title = expand_template(title_template, title_context)

    # ── Pre-process each sub-doc ───────────────────────────────
    cur = start_offset + 1
    tmp_dir = tempfile.mkdtemp(prefix='gesa_sub_')
    temp_subs = []
    for fp in file_list:
        if not os.path.exists(fp):
            continue
        sd = _open_doc(fp)
        # Fix common misspellings in text before processing
        for p in get_all_paragraphs(sd):
            for run in p.runs:
                if 'ompontencia' in run.text.lower() or 'ompeyencia' in run.text.lower():
                    run.text = re.sub(r'Compontencia|Compeyencia', 'Competencia', run.text, flags=re.IGNORECASE)
        strip_leading_empty_paras_and_breaks(sd)
        strip_section_breaks(sd)
        _resolve_autonumbering(sd)
        cur = apply_renumbering_and_ranges(sd, cur)
        process_competencias_and_componentes(sd)
        ensure_proper_spacing_between_questions(sd)
        reset_letter_sequence(sd)
        apply_formatting_to_document(sd)

        # Clear headers/footers from sub-documents so they inherit the template's
        for sec in sd.sections:
            force_single_column(sec)
            sectPr = sec._sectPr
            for ref in list(sectPr.findall(qn('w:headerReference'))):
                sectPr.remove(ref)
            for ref in list(sectPr.findall(qn('w:footerReference'))):
                sectPr.remove(ref)
            try:
                titlePg = sectPr.find(qn('w:titlePg'))
                if titlePg is not None:
                    sectPr.remove(titlePg)
            except:
                pass
            # Force vertical alignment to top
            vAlign = sectPr.find(qn('w:vAlign'))
            if vAlign is not None:
                sectPr.remove(vAlign)

        tp = os.path.join(tmp_dir, os.path.basename(fp) + '.tmp.docx')
        normalize_document_xml(sd)
        sanitize_document_xml(sd)
        sd.save(tp)
        _rebuild_zip(tp)
        temp_subs.append(tp)

    if not temp_subs:
        shutil.rmtree(tmp_dir, ignore_errors=True)
        raise RuntimeError('No hay sub-documentos para procesar.')

    # ── Pre-process template: replace placeholders everywhere ──
    tpl = _open_doc(template_path)
    try:
        import datetime
        now = datetime.datetime.utcnow()
        tpl.core_properties.created = now
        tpl.core_properties.modified = now
        tpl.core_properties.last_modified_by = "GESA"
        tpl.core_properties.author = "GESA"
        tpl.core_properties.revision = 1
    except Exception as e:
        print(f"[GESA] No se pudieron actualizar los metadatos de tpl: {e}")
    replace_in_all(tpl, replacements_map)
    for sec in tpl.sections:
        force_single_column(sec)
        # Force vertical alignment to top
        sectPr = sec._sectPr
        vAlign = sectPr.find(qn('w:vAlign'))
        if vAlign is not None:
            sectPr.remove(vAlign)

    wns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    for br in list(tpl.element.body.xpath('.//w:br[@w:type="page"]')):
        p_br = br.getparent()
        if p_br is not None:
            p_br.remove(br)
    for s in list(tpl.element.body.xpath('.//w:sectPr')):
        if s.getparent() != tpl.element.body:
            p_s = s.getparent()
            if p_s is not None:
                p_s.remove(s)

    # Trim trailing empty paragraphs in template body to max 2 spacers
    # (the template has ~8 empty ¶ that push content to page 2;
    #  2 spacers are enough to clear the floating header graphic)
    body_sect = tpl.element.body.find(f'{{{wns}}}sectPr')
    trailing_empty = []
    for elem in reversed(list(tpl.element.body)):
        if elem == body_sect:
            continue
        if elem.tag.endswith('}p') or elem.tag == 'p':
            t = ''.join(te.text or '' for te in elem.iter(f'{{{wns}}}t')).strip()
            has_draw = len(elem.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')) > 0 or len(elem.findall('.//{urn:schemas-microsoft-com:vml}shape')) > 0
            if not t and not has_draw:
                trailing_empty.append(elem)
            else:
                break
        else:
            break
    max_spacers = 2
    for emp in trailing_empty[max_spacers:]:
        tpl.element.body.remove(emp)

    _SENTINEL_TEXT = "GESABOUNDARY"
    sentinel_p = parse_xml(
        f'<w:p {nsdecls("w")}><w:r><w:rPr><w:rFonts w:ascii="Arial" w:hAnsi="Arial"/><w:sz w:val="2"/><w:szCs w:val="2"/><w:color w:val="FFFFFF"/></w:rPr><w:t xml:space="preserve">{_SENTINEL_TEXT}</w:t></w:r></w:p>'
    )
    # Insertar el sentinel ANTES del sectPr (sectPr debe ser el último hijo del body)
    body_sectPr_ = tpl.element.body.find(f'{{{wns}}}sectPr')
    if body_sectPr_ is not None:
        body_sectPr_.addprevious(sentinel_p)
    else:
        tpl.element.body.append(sentinel_p)

    prepped = os.path.join(tmp_dir, 'template_prepped.docx')
    normalize_document_xml(tpl)
    sanitize_document_xml(tpl)
    _ensure_sectPr_is_last(tpl)
    tpl.save(prepped)
    _rebuild_zip(prepped)

    # ── Word COM: merge + post-process + guardado final ────────
    # (evita el guardado de python-docx que introduce defectos XML)
    word = None
    doc = None
    try:
        import win32com.client as win32
        word = win32.DispatchEx('Word.Application')
        word.Visible = False
        word.DisplayAlerts = False
        time.sleep(0.3)

        doc = word.Documents.Open(
            os.path.abspath(prepped),
            ConfirmConversions=False, ReadOnly=False,
            AddToRecentFiles=False)
        time.sleep(0.2)

        for i, tp in enumerate(temp_subs):
            rng = doc.Range()
            rng.Collapse(0)
            rng.InsertFile(os.path.abspath(tp), ConfirmConversions=False)
            if i < len(temp_subs) - 1:
                rng = doc.Range()
                rng.Collapse(0)
                rng.InsertBreak(2)

        for s_idx in range(1, doc.Sections.Count + 1):
            try:
                doc.Sections(s_idx).PageSetup.TextColumns.SetCount(1)
                doc.Sections(s_idx).PageSetup.VerticalAlignment = 0  # wdAlignVerticalTop (0)
            except:
                pass

        # ── Unificar encabezados/pies de página de todas las secciones ──
        # InsertFile importa un salto de sección por cada subdocumento, por lo
        # que el resultado final tiene varias secciones. Cada sección debe
        # heredar el encabezado/pie de la primera; de lo contrario la
        # paginación (números de página) aparece fragmentada y distinta por
        # bloque de asignatura.
        try:
            if doc.Sections.Count > 1:
                for s_idx in range(2, doc.Sections.Count + 1):
                    try:
                        doc.Sections(s_idx).Headers(1).LinkToPrevious = True
                        doc.Sections(s_idx).Footers(1).LinkToPrevious = True
                    except Exception:
                        pass
        except Exception:
            pass

        try:
            doc.Content.Find.ClearFormatting()
            for k, v in replacements_map.items():
                doc.Content.Find.Execute(FindText=k, ReplaceWith=str(v), Replace=2)
        except:
            pass
        for i in range(1, doc.Shapes.Count + 1):
            try:
                shape = doc.Shapes.Item(i)
                if shape.Type == 6:
                    for j in range(1, shape.GroupItems.Count + 1):
                        try:
                            gi = shape.GroupItems.Item(j)
                            if gi.TextFrame.HasText:
                                gi.TextFrame.TextRange.Find.ClearFormatting()
                                for k, v in replacements_map.items():
                                    gi.TextFrame.TextRange.Find.Execute(
                                        FindText=k, ReplaceWith=str(v), Replace=2)
                        except:
                            pass
                elif shape.TextFrame.HasText:
                    shape.TextFrame.TextRange.Find.ClearFormatting()
                    for k, v in replacements_map.items():
                        shape.TextFrame.TextRange.Find.Execute(
                            FindText=k, ReplaceWith=str(v), Replace=2)
            except:
                pass

        # ── Post-process in Word COM (evita guardado de python-docx) ──
        # Forzar Century Gothic 11pt en contenido de subdocumentos
        try:
            fr = doc.Range()
            fr.Find.Text = _SENTINEL_TEXT
            if fr.Find.Execute():
                content_range = doc.Range(fr.End, doc.Content.End)
                content_range.Font.Name = "Century Gothic"
                content_range.Font.Size = 11
                fr.Text = ""
        except:
            pass
        # Pie de página con número de página
        # El template ya trae "Página X de Y" con campos PAGE/NUMPAGES; no se
        # debe insertar otro campo de número de página encima (Word lo pega al
        # texto existente y la paginación "sale rara"). Solo se agrega si el
        # pie de la primera sección no tiene ningún campo.
        try:
            if doc.Sections.Count > 0:
                ft = doc.Sections(1).Footers(1)
                ft.Range.Font.Name = "Century Gothic"
                ft.Range.Font.Size = 11
                has_field = False
                try:
                    has_field = ft.Range.Fields.Count > 0
                except Exception:
                    has_field = False
                if not has_field:
                    ft.PageNumbers.Add(PageNumberAlignment=2)
        except:
            pass
        # Ajustar márgenes de página y forzar alineación vertical superior
        try:
            ps = doc.PageSetup
            ps.TopMargin = 0.3937 * 72
            ps.BottomMargin = 0.3937 * 72
            ps.LeftMargin = 0.3937 * 72
            ps.RightMargin = 0.3937 * 72
            ps.VerticalAlignment = 0  # wdAlignVerticalTop (0)
        except:
            pass

        # Actualizar metadatos en Word COM
        try:
            doc.BuiltInDocumentProperties("Title").Value = expanded_title
            doc.BuiltInDocumentProperties("Category").Value = f"{eval_prefix} de Suficiencia Académica"
            doc.BuiltInDocumentProperties("Author").Value = "GESA"
        except Exception as e:
            print(f"[GESA] Word COM no pudo establecer propiedades: {e}")

        # Guardar a TEMPORAL LOCAL (OneDrive/cloud bloquean archivos en uso)
        tmp_word = os.path.join(tmp_dir, 'gesa_word.docx')
        doc.SaveAs2(os.path.abspath(tmp_word), FileFormat=12, AddToRecentFiles=False)
        doc.Close(False)
        doc = None
        word.Quit()
        word = None
        time.sleep(0.3)

        # ── Aplicar listas nativas al documento mergeado por Word COM ──
        try:
            final_merged = _open_doc(tmp_word)
            apply_native_lists_to_final_doc(final_merged, start_offset=start_offset)
            normalize_document_xml(final_merged)
            sanitize_document_xml(final_merged)
            _ensure_sectPr_is_last(final_merged)
            final_merged.save(tmp_word)
            _rebuild_zip(tmp_word)
            _clean_orphaned_header_footer_rels(tmp_word)
            _clean_rsid_attributes(tmp_word)

            # Re-guardar con Word COM para limpiar defectos XML introducidos por python-docx
            import win32com.client as win32
            word_clean = win32.DispatchEx('Word.Application')
            word_clean.Visible = False
            word_clean.DisplayAlerts = False
            time.sleep(0.2)
            doc_clean = word_clean.Documents.Open(
                os.path.abspath(tmp_word),
                ConfirmConversions=False, ReadOnly=False,
                AddToRecentFiles=False)
            time.sleep(0.2)
            doc_clean.Save()
            doc_clean.Close(False)
            doc_clean = None
            word_clean.Quit()
            word_clean = None
            time.sleep(0.3)
        except Exception as e_native:
            print(f"[GESA] Error aplicando listas nativas en flujo COM: {e_native}")

        # Copiar a destino (OneDrive ya no interfiere con Word COM)
        if os.path.exists(output_path):
            for _ in range(10):
                try:
                    os.remove(output_path)
                    break
                except:
                    time.sleep(0.3)
        shutil.copy2(tmp_word, output_path)
        os.remove(tmp_word)

        # Limpiar posibles archivos de bloqueo en el destino
        lock_file = os.path.join(os.path.dirname(output_path), '~$' + os.path.basename(output_path))
        if os.path.exists(lock_file):
            try: os.remove(lock_file)
            except: pass

        shutil.rmtree(tmp_dir, ignore_errors=True)
        print(f"[GESA] Word COM exitoso: {os.path.basename(output_path)}")
        return cur - 1
    except Exception as ex:
        print(f"[GESA] Word COM fallo en merge: {ex}")
    finally:
        if doc is not None:
            try: doc.Close(False)
            except: pass
        if word is not None:
            try: word.Quit()
            except: pass
        doc = None
        word = None

    # ── Fallback: python-docx merge + post-process ────────────
    print(f"[GESA] Word COM no disponible, usando python-docx: {os.path.basename(output_path)}")
    merged_path = os.path.join(tmp_dir, 'word_merged.docx')

    final_doc = _open_doc(prepped)
    for i, tp in enumerate(temp_subs):
        if stop_check is not None and stop_check():
            raise RuntimeError("Cancelado por el usuario.")
        sub_doc = _open_doc(tp)
        _merge_docx_with_rels(final_doc, sub_doc, add_break=(i < len(temp_subs) - 1))
    normalize_document_xml(final_doc)
    sanitize_document_xml(final_doc)
    _ensure_sectPr_is_last(final_doc)
    final_doc.save(merged_path)
    _rebuild_zip(merged_path)

    # ── Post-process (python-docx) ──────────────────────────────
    final = _open_doc(merged_path)
    strip_section_breaks(final)
    process_competencias_and_componentes(final)
    process_habilidades_and_bullets(final)
    ensure_proper_spacing_between_questions(final)
    expanded_title = expand_template(title_template, title_context)
    final.core_properties.title = expanded_title
    final.core_properties.category = f"{eval_prefix} de Suficiencia Académica"
    final.core_properties.content_status = config_data.get('period', '')
    try:
        import datetime
        now = datetime.datetime.utcnow()
        final.core_properties.created = now
        final.core_properties.modified = now
        final.core_properties.last_modified_by = "GESA"
        final.core_properties.author = "GESA"
        final.core_properties.revision = 1
    except Exception as e:
        print(f"[GESA] No se pudieron actualizar los metadatos de final: {e}")

    for idx, sec in enumerate(final.sections):
        if idx == 0:
            apply_section0_page_setup(sec)
        else:
            apply_subsequent_page_setup(sec)

        # Strip all header references from every section (defensive)
        sectPr = sec._sectPr
        for ref in list(sectPr.findall(qn('w:headerReference'))):
            sectPr.remove(ref)
        titlePg = sectPr.find(qn('w:titlePg'))
        if titlePg is not None:
            sectPr.remove(titlePg)
        # Footer page number
        sec.footer.is_linked_to_previous = False
        setup_footer_page_number(sec.footer, doc=final)

    # ── Force Century Gothic 11pt on subdoc content (after sentinel) ──
    # The template body content (before the sentinel) keeps its original font.
    wns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    font_name = "Century Gothic"
    _SENTINEL_TEXT = "GESABOUNDARY"

    def _force_run_font(r_elem):
        rPr = r_elem.find(f'{{{wns}}}rPr')
        if rPr is None:
            rPr = parse_xml(f'<w:rPr {nsdecls("w")}/>')
            r_elem.insert(0, rPr)
        rFonts = rPr.find(f'{{{wns}}}rFonts')
        if rFonts is None:
            rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{font_name}" w:hAnsi="{font_name}" w:cs="{font_name}"/>')
            rPr.append(rFonts)
        else:
            ascii_f = (rFonts.get(qn('w:ascii')) or '').lower()
            if not any(sym in ascii_f for sym in ['symbol', 'wingdings', 'webdings', 'marlett']):
                rFonts.set(qn('w:ascii'), font_name)
                rFonts.set(qn('w:hAnsi'), font_name)
                rFonts.set(qn('w:cs'), font_name)
                for theme_attr in ['w:asciiTheme', 'w:hAnsiTheme', 'w:eastAsiaTheme', 'w:csTheme', 'w:theme']:
                    if qn(theme_attr) in rFonts.attrib:
                        del rFonts.attrib[qn(theme_attr)]
        sz = rPr.find(f'{{{wns}}}sz')
        if sz is None:
            sz = parse_xml(f'<w:sz {nsdecls("w")} w:val="22"/>')
            rPr.append(sz)
        else:
            sz.set(qn('w:val'), '22')
        szCs = rPr.find(f'{{{wns}}}szCs')
        if szCs is None:
            szCs = parse_xml(f'<w:szCs {nsdecls("w")} w:val="22"/>')
            rPr.append(szCs)
        else:
            szCs.set(qn('w:val'), '22')

    # Locate sentinel by text content
    body_children = list(final.element.body)
    sentinel_idx = None
    for idx_c, child in enumerate(body_children):
        if child.tag.endswith('}p') or child.tag == 'p':
            for t in child.xpath('.//w:t'):
                if _SENTINEL_TEXT in (t.text or ''):
                    sentinel_idx = idx_c
                    break
        if sentinel_idx is not None:
            break

    # Strip page breaks only from sentinel paragraph if any
    if sentinel_idx is not None:
        child = body_children[sentinel_idx]
        if child.tag.endswith('}p') or child.tag == 'p':
            for br in list(child.xpath('.//w:br[@w:type="page"]')):
                p_br = br.getparent()
                if p_br is not None:
                    p_br.remove(br)

    if sentinel_idx is not None:
        try:
            s_elem = body_children[sentinel_idx]
            for t in s_elem.xpath('.//w:t'):
                t.text = (t.text or '').replace(_SENTINEL_TEXT, '')
            s_elem.getparent().remove(s_elem)
        except Exception:
            pass
        body_children = list(final.element.body)
        subdoc_elems = body_children[sentinel_idx:]
    else:
        subdoc_elems = []

    # Safety: remove any leftover fragments containing sentinel text
    for t in final.element.body.xpath(f'.//w:t[contains(text(), "{_SENTINEL_TEXT}")]'):
        t.text = (t.text or '').replace(_SENTINEL_TEXT, '')
    for p_elem in final.element.body.xpath('.//w:p'):
        text_content = ''.join(t.text or '' for t in p_elem.xpath('.//w:t'))
        if _SENTINEL_TEXT in text_content:
            try:
                t_elem = p_elem.xpath(f'.//w:t[contains(text(), "{_SENTINEL_TEXT}")]')
                if t_elem:
                    t_elem[0].text = (t_elem[0].text or '').replace(_SENTINEL_TEXT, '')
            except Exception:
                pass

    from docx.text.paragraph import Paragraph as _Paragraph
    for elem in subdoc_elems:
        if elem.tag.endswith('}p') or elem.tag == 'p':
            para = _Paragraph(elem, final)
            set_single_line_spacing(para)
            text = para.text.strip()

            is_comp = bool(re.match(r'^(Competencia|Competence|Componente|Component|Habilidades|Habilidad|Nivel|Level|Desempeño|Aprendizaje|Afirmación|Estándar)\s*[:\-]', text, re.IGNORECASE))
            is_q_or_opt = bool(re.match(r'^(\s*[\(\[\{]?(\d+|[a-eA-E])\s*[\.\)\]\}\-\:\/])', text))
            is_bul = is_bullet_paragraph(para)
            if not _has_drawing(elem):
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT if (is_comp or is_q_or_opt or is_bul) else WD_ALIGN_PARAGRAPH.JUSTIFY
        for r_elem in elem.iter(f'{{{wns}}}r'):
            _force_run_font(r_elem)

    _fix_numbering_level_fonts(final, font_name)

    # Headers — force Century Gothic 11pt
    for sec in final.sections:
        for hdr in (sec.header, sec.first_page_header):
            if hdr is not None:
                for r_elem in hdr._element.xpath('.//w:r'):
                    _force_run_font(r_elem)

    # Footer pagination — force Century Gothic 11pt
    for sec in final.sections:
        if sec.footer is not None:
            for r_elem in sec.footer._element.xpath('.//w:r'):
                _force_run_font(r_elem)

    # Apply native numbering to the fully merged document
    apply_native_lists_to_final_doc(final, start_offset=start_offset)
    normalize_document_xml(final)

    # Sanitizar caracteres inválidos para XML antes de guardar
    sanitize_document_xml(final)

    _ensure_sectPr_is_last(final)

    if os.path.exists(output_path):
        for _ in range(5):
            try:
                os.remove(output_path)
                break
            except Exception:
                time.sleep(0.3)

    try:
        final.save(output_path)
    except PermissionError:
        raise PermissionError(f"El archivo '{os.path.basename(output_path)}' est\u00e1 abierto en Microsoft Word. Por favor ci\u00e9rralo e intenta de nuevo.")
    except Exception as err:
        raise RuntimeError(f"No se pudo guardar '{os.path.basename(output_path)}': {err}")

    # Reconstruir ZIP para eliminar defectos de compresión/estructura
    _rebuild_zip(output_path)
    # Limpiar relaciones huérfanas de headers/footers que Word COM rechaza
    _clean_orphaned_header_footer_rels(output_path)
    # Limpiar atributos rsidR/rsidP que python-docx agrega y Word rechaza
    _clean_rsid_attributes(output_path)

    # ── Re-guardar con Word COM para limpiar defectos XML de python-docx ──
    word_clean = None
    doc_clean = None
    try:
        import win32com.client as win32
        word_clean = win32.DispatchEx('Word.Application')
        word_clean.Visible = False
        word_clean.DisplayAlerts = False
        import time as _time
        _time.sleep(0.2)
        doc_clean = word_clean.Documents.Open(
            os.path.abspath(output_path),
            ConfirmConversions=False, ReadOnly=False,
            AddToRecentFiles=False)
        _time.sleep(0.2)
        doc_clean.SaveAs2(
            os.path.abspath(output_path),
            FileFormat=12,
            AddToRecentFiles=False)
        doc_clean.Close(False)
        doc_clean = None
        word_clean.Quit()
        word_clean = None
        print(f"[GESA] Word COM re-save OK: {os.path.basename(output_path)}")
    except Exception as _e:
        print(f"[GESA] Word COM re-save: se conserva version python-docx ({_e})")
    finally:
        if doc_clean is not None:
            try: doc_clean.Close(False)
            except: pass
        if word_clean is not None:
            try: word_clean.Quit()
            except: pass
        doc_clean = None
        word_clean = None

    shutil.rmtree(tmp_dir, ignore_errors=True)
    return cur - 1
